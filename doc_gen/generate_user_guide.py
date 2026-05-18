"""
generate_user_guide.py
----------------------
Generates the doc-gen User Guide as a Word document.
Run: python generate_user_guide.py
Output: ../docs/doc_gen_user_guide.docx
"""

from __future__ import annotations
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ────────────────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x1F, 0x4E, 0x79)
C_MID    = RGBColor(0x2E, 0x75, 0xB6)
C_GREEN  = RGBColor(0x37, 0x86, 0x30)
C_ORANGE = RGBColor(0xC5, 0x5A, 0x11)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_GRAY   = RGBColor(0x80, 0x80, 0x80)

HDR_FILL = "D5E8F0"
ALT_FILL = "EBF3FB"
GRN_FILL = "E2EFDA"
ORG_FILL = "FCE4D6"
COD_FILL = "F2F2F2"
ACC_COLOR = "2E75B6"
GRY_COLOR = "BFBFBF"

FONT = "微軟正黑體"
MONO = "Courier New"

OUTPUT = str(Path(__file__).parent.parent / "docs" / "doc_gen_user_guide.docx")


# ── XML helpers ────────────────────────────────────────────────────────────────
def _cell_bg(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _cell_borders(cell, color=GRY_COLOR):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'), color)
        tcB.append(el)
    tcPr.append(tcB)

def _para_border_bottom(p, color=ACC_COLOR, sz=8):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz))
    b.set(qn('w:space'),'4'); b.set(qn('w:color'), color)
    pBdr.append(b); pPr.append(pBdr)

def _col_w(table, ci, cm):
    dxa = int(cm * 567)
    for row in table.rows:
        tc = row.cells[ci]._tc; tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(dxa)); tcW.set(qn('w:type'),'dxa')
        tcPr.append(tcW)

def _page_break(doc):
    p = doc.add_paragraph(); r = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'),'page'); r._r.append(br)

def _add_header_footer(doc, title):
    sec = doc.sections[0]
    # header
    hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(title); r.font.size=Pt(9); r.font.name=FONT; r.font.color.rgb=C_GRAY
    _para_border_bottom(hp, GRY_COLOR, 4)
    # footer
    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt, fld in [("第 ",None),(None,"PAGE"),(" / ",None),(None,"NUMPAGES"),(" 頁",None)]:
        if txt:
            rr = fp.add_run(txt)
        else:
            rr = fp.add_run()
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'),'begin'); rr._r.append(fc)
            it = OxmlElement('w:instrText'); it.text=f' {fld} '; rr._r.append(it)
            rr2 = fp.add_run()
            fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'),'end'); rr2._r.append(fc2)
            rr = rr2
        rr.font.size=Pt(9); rr.font.name=FONT; rr.font.color.rgb=C_GRAY


# ── Builder helpers ────────────────────────────────────────────────────────────
class B:
    def __init__(self):
        self.doc = Document()
        sec = self.doc.sections[0]
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.0)
        self.doc.styles['Normal'].font.name = FONT
        self.doc.styles['Normal'].font.size = Pt(10.5)

    def cover(self, main, sub, meta: list[tuple]):
        doc = self.doc
        doc.add_paragraph().paragraph_format.space_before = Pt(60)
        for txt, sz, bold, color in [
            (main, 28, True, C_DARK),
            (sub,  18, False, C_MID),
        ]:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(txt); r.font.size=Pt(sz); r.font.bold=bold
            r.font.name=FONT; r.font.color.rgb=color
            p.paragraph_format.space_after = Pt(4)
        hr = doc.add_paragraph(); _para_border_bottom(hr, ACC_COLOR, 12)
        doc.add_paragraph().paragraph_format.space_before = Pt(10)
        for label, val in meta:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p.add_run(f"{label}："); r1.font.bold=True; r1.font.size=Pt(11); r1.font.name=FONT
            r2 = p.add_run(val); r2.font.size=Pt(11); r2.font.name=FONT
            p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)

    def h1(self, txt):
        p = self.doc.add_paragraph(style='Heading 1'); p.clear()
        p.paragraph_format.space_before=Pt(18); p.paragraph_format.space_after=Pt(6)
        r = p.add_run(txt); r.font.size=Pt(16); r.font.bold=True
        r.font.name=FONT; r.font.color.rgb=C_DARK
        _para_border_bottom(p)

    def h2(self, txt):
        p = self.doc.add_paragraph(style='Heading 2'); p.clear()
        p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
        r = p.add_run(txt); r.font.size=Pt(13); r.font.bold=True
        r.font.name=FONT; r.font.color.rgb=C_MID

    def h3(self, txt):
        p = self.doc.add_paragraph(style='Heading 3'); p.clear()
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
        r = p.add_run(txt); r.font.size=Pt(11); r.font.bold=True
        r.font.name=FONT; r.font.color.rgb=C_DARK

    def body(self, txt, bold=False, color=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r = p.add_run(txt); r.font.size=Pt(10.5); r.font.bold=bold; r.font.name=FONT
        if color: r.font.color.rgb=color
        return p

    def bullet(self, txt, bold_prefix=None):
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
        if bold_prefix:
            r1 = p.add_run(bold_prefix); r1.font.bold=True; r1.font.name=FONT; r1.font.size=Pt(10.5)
        r2 = p.add_run(txt); r2.font.name=FONT; r2.font.size=Pt(10.5)

    def numbered(self, txt):
        p = self.doc.add_paragraph(style='List Number')
        r = p.add_run(txt); r.font.name=FONT; r.font.size=Pt(10.5)

    def code(self, lines: list[str], caption=None):
        if caption:
            pc = self.doc.add_paragraph()
            pc.paragraph_format.space_before=Pt(6)
            rc = pc.add_run(caption); rc.font.bold=True; rc.font.size=Pt(9.5); rc.font.name=FONT; rc.font.color.rgb=C_GRAY
        for line in lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            p.paragraph_format.left_indent=Cm(0.8)
            r = p.add_run(line); r.font.name=MONO; r.font.size=Pt(9)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),COD_FILL)
            pPr.append(shd)

    def callout(self, icon, label, txt, fill, text_color=None):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent=Cm(0.5)
        p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
        r1 = p.add_run(f"{icon} {label}："); r1.font.bold=True; r1.font.name=FONT; r1.font.size=Pt(10.5)
        if text_color: r1.font.color.rgb=text_color
        r2 = p.add_run(txt); r2.font.name=FONT; r2.font.size=Pt(10.5)
        if text_color: r2.font.color.rgb=text_color
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), fill)
        pPr.append(shd)

    def note(self, txt):
        self.callout("!", "注意", txt, ORG_FILL, C_ORANGE)

    def tip(self, txt):
        self.callout("*", "提示", txt, GRN_FILL, C_GREEN)

    def spacer(self, pts=6):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before=Pt(pts); p.paragraph_format.space_after=Pt(0)

    def table(self, headers, rows, widths_cm, header_fill=HDR_FILL):
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
        for ci,w in enumerate(widths_cm): _col_w(tbl, ci, w)
        # header
        hr = tbl.rows[0]
        for ci,h in enumerate(headers):
            cell=hr.cells[ci]; cell.text=''
            _cell_bg(cell, header_fill); _cell_borders(cell, ACC_COLOR)
            cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            p=cell.paragraphs[0]; r=p.add_run(h)
            r.font.bold=True; r.font.name=FONT; r.font.size=Pt(10); r.font.color.rgb=C_DARK
        # rows
        for ri,row in enumerate(rows):
            tr=tbl.add_row(); alt=(ri%2==1)
            for ci,val in enumerate(row):
                cell=tr.cells[ci]; cell.text=''
                if alt: _cell_bg(cell, ALT_FILL)
                _cell_borders(cell)
                cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                p=cell.paragraphs[0]; r=p.add_run(str(val) if val else '—')
                r.font.name=FONT; r.font.size=Pt(10)
        return tbl

    def inline_code(self, txt):
        """Return a run styled as inline code (used inside body paragraphs)."""
        return txt  # caller handles runs manually

    def save(self, path):
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)
        print(f"Saved: {path}  ({Path(path).stat().st_size//1024} KB)")


# ── Build guide ────────────────────────────────────────────────────────────────
def build():
    b = B()
    _add_header_footer(b.doc, "doc-gen 使用者指南 — 台灣自來水公司 資訊整合系統")

    # ══ COVER ══════════════════════════════════════════════════════════════════
    b.cover(
        "doc-gen 自動文件產生工具",
        "使用者指南 (User Guide)",
        [
            ("版本",    "1.0"),
            ("日期",    datetime.now().strftime("%Y-%m-%d")),
            ("適用專案","台灣自來水公司 資訊整合系統"),
            ("工具語言","Python 3.9+"),
        ]
    )
    _page_break(b.doc)

    # ══ 1. 概述 ════════════════════════════════════════════════════════════════
    b.h1("1. 工具概述")

    b.h2("1.1 這個工具是什麼？")
    b.body(
        "doc-gen 是一個 Pure Python 程式，能夠自動掃描 Java / JSP / Properties 原始碼，"
        "分析專案架構，並產生格式完整的 Word 設計文件（.docx）。"
    )
    b.body("每天早上 10:00，Windows 工作排程器會自動觸發此程式，無需任何人工介入。")

    b.spacer(4)
    b.h2("1.2 它使用 Claude Code 或 AI 嗎？")
    b.callout(
        "X", "不使用 AI",
        "doc-gen 完全不依賴 Claude Code、OpenAI、任何 LLM API 或網際網路連線。"
        "它是 100% 離線的靜態程式碼分析工具，所有邏輯皆由 Python 標準程式庫與正規表達式實作。",
        ORG_FILL, C_ORANGE
    )
    b.spacer(2)
    b.table(
        ["比較項目", "doc-gen（本工具）", "Claude Code / AI 工具"],
        [
            ["運作方式",   "正規表達式靜態掃描原始碼",   "Large Language Model 推論"],
            ["網路需求",   "完全離線",                  "需要 API 連線"],
            ["速度",       "< 5 秒",                   "依 API 回應時間，通常 10～60 秒"],
            ["費用",       "免費（Python 開源套件）",    "依 token 用量計費"],
            ["準確性",     "精確反映程式碼現況",          "可能有幻覺（hallucination）"],
            ["隱私安全",   "程式碼不離開本機",            "程式碼傳送至雲端"],
        ],
        [4.0, 5.5, 6.5]
    )

    b.spacer(4)
    b.h2("1.3 工具架構總覽")
    b.body("doc-gen 由三個 Python 模組組成，以 Windows Task Scheduler 觸發：")
    b.spacer(2)
    b.code([
        "doc_gen/",
        "  analyzer.py          ← 原始碼解析引擎（靜態分析）",
        "  word_builder.py      ← Word 文件建構器",
        "  doc_generator.py     ← 主程式入口 + 排程 Daemon 模式",
        "  requirements.txt     ← 相依套件清單（只有 python-docx）",
        "  setup_scheduler.bat  ← Windows 排程一鍵設定（雙擊執行）",
        "  setup_scheduler.ps1  ← Windows 排程設定（PowerShell 版）",
    ], caption="檔案結構")

    _page_break(b.doc)

    # ══ 2. 分析工具說明 ═════════════════════════════════════════════════════════
    b.h1("2. 分析工具詳解")

    b.h2("2.1 使用哪些 Python 套件？")
    b.table(
        ["套件 / 模組", "類型", "用途", "安裝方式"],
        [
            ["re",           "Python 標準函式庫", "正規表達式，用於解析 Java/JSP 原始碼",     "內建，無需安裝"],
            ["pathlib",      "Python 標準函式庫", "跨平台檔案路徑操作、遞迴掃描目錄",           "內建，無需安裝"],
            ["dataclasses",  "Python 標準函式庫", "定義資料模型（JavaClass、Module 等結構）",  "內建，無需安裝"],
            ["argparse",     "Python 標準函式庫", "解析命令列參數（--once, --daemon, --root）","內建，無需安裝"],
            ["logging",      "Python 標準函式庫", "結構化日誌（同時輸出 console + 檔案）",      "內建，無需安裝"],
            ["time / datetime","Python 標準函式庫","排程計時、日期格式化（文件檔名）",           "內建，無需安裝"],
            ["python-docx",  "第三方套件",        "建立 Word .docx 文件（表格、標題、頁首頁尾）","pip install python-docx"],
        ],
        [3.5, 3.0, 6.5, 3.0]
    )
    b.tip("全部相依只有一個外部套件：python-docx。執行 pip install python-docx 即可完成安裝。")

    b.spacer(4)
    b.h2("2.2 analyzer.py — 靜態程式碼分析引擎")
    b.body("analyzer.py 使用 Python re（正規表達式）對原始碼進行純文字解析，不需要編譯器或 JVM。")
    b.body("以下是它偵測的每一種資訊及對應的正規表達式：", bold=True)
    b.spacer(2)
    b.table(
        ["偵測目標", "正規表達式模式（示意）", "結果欄位"],
        [
            ["Java 套件名稱",      r"package\s+([\w.]+)\s*;",                            "JavaClass.package"],
            ["類別 / 介面 / 列舉", r"(class|interface|enum)\s+(\w+)",                    "JavaClass.name, kind"],
            ["繼承 / 實作",        r"extends\s+([\w]+)|implements\s+([\w, ]+)",           "JavaClass.parent"],
            ["公開方法",           r"(public|protected)\s+[\w<>]+\s+(\w+)\s*\(",         "JavaClass.methods"],
            ["@Api 端點路徑",      r'@Api\([^)]*path\s*=\s*"([^"]+)"',                  "JavaClass.api_paths"],
            ["SQL 資料表名稱",     r'(?:FROM|JOIN|INTO|UPDATE)\s+([A-Za-z_]\w+)',        "JavaClass.db_tables"],
            ["Unit ID 常數",       r'UNIT_ID\s*=\s*"([^"]+)"',                           "JavaClass.unit_id"],
            ["Cron 排程表達式",    r'(?:cron|schedule)\s*=\s*"([^"]+)"',                 "JavaClass.timer_schedule"],
            ["頂層 Annotation",    r'@(\w+)(?:\([^)]*\))?',                              "JavaClass.annotations"],
        ],
        [4.0, 7.0, 5.0]
    )

    b.spacer(4)
    b.h3("2.2.1 掃描流程（三步驟）")
    b.numbered("目錄遍歷：使用 pathlib.Path.rglob() 遞迴掃描專案根目錄下所有 .java、.jsp、.properties 檔案。")
    b.numbered("過濾雜訊：自動跳過備份資料夾（ciopy、複製、archive、update）、暫存目錄（.tmp_ccms）、建置輸出（target、build）。")
    b.numbered("模組歸屬：依路徑判斷檔案屬於哪個模組（ecpsso / waterCusSSO / waterHrSync）。")

    b.spacer(4)
    b.h3("2.2.2 過濾規則")
    b.table(
        ["過濾條件", "範例", "原因"],
        [
            ["資料夾名稱含 ciopy",            "waterImportTimer - ciopy/", "複製品（ciopy = 複製 Pinyin），內容與 waterHrSync 完全相同"],
            ["資料夾名稱含 複製 / copy / bak", "sso - 複製.jsp",           "手動備份檔案，不是正式原始碼"],
            ["update / archive 資料夾",       "update/archive/202604.jsp", "歷史版本歸檔，非當前版本"],
            [".tmp_ccms 資料夾",              ".tmp_ccms/autoLogin.jsp",   "IDE 或建置工具產生的暫存檔"],
            ["target / build 資料夾",         "target/classes/",           "Maven 編譯輸出，不是原始碼"],
        ],
        [5.0, 5.0, 6.0]
    )
    b.note("過濾規則定義在 analyzer.py 的 SKIP_DIRS 與 COPY_FOLDER_MARKERS 常數中，可依需求修改。")

    b.spacer(4)
    b.h2("2.3 word_builder.py — Word 文件建構器")
    b.body("word_builder.py 使用 python-docx 套件，將 analyzer.py 產生的資料模型轉換為格式完整的 Word 文件。")
    b.spacer(2)
    b.table(
        ["功能", "實作方式", "python-docx API"],
        [
            ["封面頁",          "置中大字標題 + 水平分隔線 + 版本資訊",          "Paragraph, Run, border XML"],
            ["章節標題",        "套用 Heading 1/2/3 樣式 + 自訂藍色底線",        "add_paragraph(style='Heading 1')"],
            ["資料表格",        "交錯列底色（藍/白）+ 深藍色標題列",              "add_table(), cell shading XML"],
            ["程式碼區塊",      "Courier New 字型 + 灰色背景",                   "Run(font=Courier New), shading XML"],
            ["ASCII 架構圖",    "逐行程式碼區塊，模擬方塊圖",                     "code block paragraphs"],
            ["警告 / 提示框",   "橘色（警告）/ 綠色（提示）底色文字段落",         "shading XML on paragraph"],
            ["頁首 / 頁尾",     "頁首右對齊文件標題；頁尾置中顯示頁碼",           "Header, Footer, PAGE field"],
            ["原子性寫檔",      "先寫入 .tmp.docx，成功後重新命名",               "Path.rename() after save()"],
        ],
        [3.5, 6.5, 6.0]
    )

    _page_break(b.doc)

    # ══ 3. 安裝 ════════════════════════════════════════════════════════════════
    b.h1("3. 安裝與環境需求")

    b.h2("3.1 系統需求")
    b.table(
        ["項目", "需求", "備註"],
        [
            ["作業系統",  "Windows 10 / 11 或 Windows Server 2019+", "Linux/macOS 亦可運行，但排程需改用 cron"],
            ["Python",    "3.9 以上",                                "建議 3.12（目前測試版本）"],
            ["磁碟空間",  "< 10 MB（程式）+ 每份文件約 50 KB",        "保留 30 份約需 1.5 MB"],
            ["網路",      "不需要",                                   "完全離線運作"],
            ["Word",      "選用（開啟文件用）",                        "Microsoft Word 或 LibreOffice Writer"],
        ],
        [3.0, 6.5, 6.5]
    )

    b.spacer(4)
    b.h2("3.2 安裝步驟")
    b.numbered("確認 Python 已安裝並在 PATH 中：")
    b.code(["python --version", "# 應顯示 Python 3.9.x 以上"])
    b.numbered("安裝唯一的外部相依套件：")
    b.code([
        "cd doc_gen",
        "pip install -r requirements.txt",
        "# requirements.txt 內容：python-docx>=1.1.0",
    ])
    b.numbered("確認安裝成功：")
    b.code([
        "python -c \"import docx; print('python-docx OK, version:', docx.__version__)\"",
    ])
    b.numbered("設定環境變數（選用，也可使用命令列參數）：")
    b.code([
        "# 指定專案根目錄（預設為 doc_gen/ 的上上層目錄）",
        "set DOC_GEN_PROJECT_ROOT=D:\\project\\Ai3\\台水\\water",
        "",
        "# 指定輸出目錄（預設為 docs/generated/）",
        "set DOC_GEN_OUTPUT_DIR=D:\\output\\docs",
        "",
        "# 指定每日執行時間（預設 10:00）",
        "set DOC_GEN_TIME=08:30",
        "",
        "# 保留最近幾份文件（預設 30，0 = 永久保留）",
        "set DOC_GEN_KEEP_LAST=30",
    ], caption="環境變數設定（在 cmd.exe 中）")

    _page_break(b.doc)

    # ══ 4. 使用方式 ════════════════════════════════════════════════════════════
    b.h1("4. 使用方式")

    b.h2("4.1 立即產生一次（手動執行）")
    b.body("最簡單的使用方式，適合測試或臨時需要最新文件時：")
    b.code([
        "cd D:\\project\\Ai3\\台水\\water\\.claude\\worktrees\\zealous-moser-336b46\\doc_gen",
        "python doc_generator.py --once",
        "",
        "# 執行結果範例：",
        "# 2026-05-18 10:00:01 [INFO] Analysing project: D:\\...",
        "# 2026-05-18 10:00:01 [INFO] Found 3 modules, 28 classes, 3 JSP files.",
        "# 2026-05-18 10:00:02 [INFO] Building Word document ...",
        "# 2026-05-18 10:00:06 [INFO] Document saved: ...\\台水資訊整合系統_設計文件_20260518.docx (50 KB)",
        "# [OK] Document ready: D:\\...\\docs\\generated\\台水資訊整合系統_設計文件_20260518.docx",
    ], caption="手動執行指令")

    b.spacer(4)
    b.h2("4.2 所有命令列參數")
    b.table(
        ["參數", "說明", "預設值", "範例"],
        [
            ["--once",   "立即執行一次後結束",                         "（無旗標時為預設行為）",     "python doc_generator.py --once"],
            ["--daemon", "持續執行，每天在 --time 指定時間自動產生",   "—",                       "python doc_generator.py --daemon"],
            ["--root",   "指定要分析的專案根目錄路徑",                 "doc_gen/ 的上上層目錄",    '--root "D:\\myproject"'],
            ["--out",    "指定輸出目錄",                               "docs/generated/",         '--out "D:\\output"'],
            ["--time",   "Daemon 模式下每日執行時間（HH:MM）",         "10:00",                   "--time 08:30"],
        ],
        [2.0, 5.5, 4.5, 4.0]
    )

    b.spacer(4)
    b.h2("4.3 設定 Windows 每日自動排程")

    b.h3("方法 A：雙擊 setup_scheduler.bat（最簡單）")
    b.numbered("在檔案總管中，進入 doc_gen 資料夾")
    b.numbered("雙擊 setup_scheduler.bat")
    b.numbered("視窗會自動偵測 Python、安裝套件、建立排程，並詢問是否立即測試執行一次")
    b.numbered("看到「Task registered OK」即完成設定")

    b.spacer(2)
    b.h3("方法 B：PowerShell 執行（進階）")
    b.code([
        "# 在任意 PowerShell 視窗執行：",
        "powershell -ExecutionPolicy Bypass -File doc_gen\\setup_scheduler.ps1",
    ])

    b.spacer(2)
    b.h3("方法 C：完全手動（schtasks）")
    b.code([
        "schtasks /Create ^",
        "  /TN \"TaiwanWater_DocGen\" ^",
        "  /TR \"python \\\"doc_gen\\doc_generator.py\\\" --once\" ^",
        "  /SC DAILY /ST 10:00 /F",
    ])

    b.spacer(2)
    b.note("排程任務名稱為 TaiwanWater_DocGen（純 ASCII），避免 cmd.exe 中文編碼問題。")

    b.spacer(4)
    b.h2("4.4 管理排程任務")
    b.table(
        ["操作", "PowerShell 指令"],
        [
            ["立即手動觸發",   "Start-ScheduledTask -TaskName 'TaiwanWater_DocGen'"],
            ["查看上次執行結果", "Get-ScheduledTask 'TaiwanWater_DocGen' | Get-ScheduledTaskInfo | Select LastRunTime,LastTaskResult"],
            ["查看排程設定",   "Get-ScheduledTask -TaskName 'TaiwanWater_DocGen' | Format-List *"],
            ["暫停排程",       "Disable-ScheduledTask -TaskName 'TaiwanWater_DocGen'"],
            ["恢復排程",       "Enable-ScheduledTask  -TaskName 'TaiwanWater_DocGen'"],
            ["刪除排程",       "Unregister-ScheduledTask -TaskName 'TaiwanWater_DocGen' -Confirm:$false"],
        ],
        [4.5, 11.5]
    )

    _page_break(b.doc)

    # ══ 5. 輸出文件說明 ═════════════════════════════════════════════════════════
    b.h1("5. 輸出文件說明")

    b.h2("5.1 輸出位置與檔名")
    b.code([
        "# 預設輸出路徑：",
        "<project_root>\\docs\\generated\\台水資訊整合系統_設計文件_YYYYMMDD.docx",
        "",
        "# 範例：",
        "docs\\generated\\台水資訊整合系統_設計文件_20260518.docx",
        "docs\\generated\\台水資訊整合系統_設計文件_20260519.docx",
        "docs\\generated\\台水資訊整合系統_設計文件_20260520.docx",
        "...",
        "# 自動保留最近 30 份，更舊的自動刪除",
    ])

    b.spacer(4)
    b.h2("5.2 文件內容結構")
    b.table(
        ["章節", "內容", "資料來源"],
        [
            ["封面頁",                   "文件標題、產生時間、版本",               "執行時自動填入"],
            ["第一部分：架構圖",          "6 張 ASCII 架構圖",                     "analyzer.py 分析結果"],
            ["  1. 系統整體架構",         "系統邊界、外部系統、使用者",              "靜態（根據專案固定描述）"],
            ["  2. 容器 / 模組架構",      "各模組（JAR/WAR）及其檔案數",            "掃描到的模組數量與 Unit ID"],
            ["  3. SSO 登入流程",         "帳密登入 → Token 驗證 → 稽核時序圖",    "靜態（根據已知業務邏輯）"],
            ["  4. HR 同步流程",          "每日凌晨同步十步驟",                     "靜態"],
            ["  5. 資料庫 ER 圖",         "資料表關聯 + 自動偵測到的所有資料表",    "SQL FROM/JOIN 正規表達式偵測"],
            ["  6. 部署架構圖",           "JBoss、DB 伺服器、AD 拓撲",              "靜態"],
            ["第二部分：詳細設計文件",    "10 章完整設計",                          "—"],
            ["  1. 專案概述",             "系統邊界表格、使用者",                    "靜態"],
            ["  2. 技術堆疊",             "Java 版本、框架、套件",                   "靜態"],
            ["  3. 模組說明",             "每個模組的類別表格（層次/方法/API）",     "掃描 .java 檔案"],
            ["    ↳ Java 類別",           "Home / Model / DAO / Service 各層類別",  "regex 解析 class 宣告"],
            ["    ↳ @Api 端點",           "偵測到的 REST API 路徑",                 "regex 解析 @Api 標註"],
            ["    ↳ Timer 排程",          "Cron 表達式",                            "regex 解析 cron/schedule 常數"],
            ["    ↳ JSP 端點",            "HTTP 方法、CORS、Request 欄位",          "掃描 .jsp 檔案"],
            ["    ↳ 設定檔",              "properties 檔案的 key=value（密碼遮蔽）", "掃描 .properties 檔案"],
            ["  4. 資料庫設計",           "自訂資料表欄位、ECP 核心表格",            "靜態（手動維護）"],
            ["  5. API 規格",             "Active / Disabled endpoints",             "靜態"],
            ["  6. 排程設計",             "Timer class 執行時間",                    "靜態"],
            ["  7. 安全設計",             "Token 機制、密碼安全、CORS",              "靜態"],
            ["  8. 錯誤處理",             "各情境處理方式",                          "靜態"],
            ["  9. 部署與設定",           "部署步驟、properties 設定、環境需求",     "靜態"],
            ["  10. 已知限制",            "已知問題 + 改善建議",                     "靜態"],
        ],
        [4.5, 6.5, 5.0]
    )

    b.spacer(4)
    b.h2("5.3 動態 vs 靜態內容")
    b.body("並非所有內容都是動態生成的。下表說明哪些章節會隨原始碼變更自動更新：")
    b.spacer(2)
    b.table(
        ["內容類型", "說明", "範例"],
        [
            ["動態（自動更新）", "每次執行重新掃描原始碼產生", "模組清單、類別名稱、方法、@Api 路徑、SQL 資料表、Timer Cron、Properties 值"],
            ["靜態（手動維護）", "固定內容，程式碼維護在 word_builder.py", "架構圖 ASCII 文字、安全設計、已知限制、部署步驟"],
        ],
        [3.5, 5.5, 7.0],
        header_fill=HDR_FILL
    )
    b.tip("若需更新靜態章節（例如新增「已知限制」項目），請直接修改 word_builder.py 對應的 rows 陣列。")

    _page_break(b.doc)

    # ══ 6. 設定與客製化 ═════════════════════════════════════════════════════════
    b.h1("6. 設定與客製化")

    b.h2("6.1 設定參數一覽（doc_generator.py 頂部）")
    b.code([
        "# 分析的專案根目錄（也可用 --root 參數或環境變數覆蓋）",
        "DEFAULT_PROJECT_ROOT = r'D:\\project\\Ai3\\台水\\water\\...'",
        "",
        "# 輸出目錄",
        "DEFAULT_OUTPUT_DIR = '...\\docs\\generated'",
        "",
        "# 輸出檔名模板（{date} 替換為 YYYYMMDD）",
        "FILENAME_TEMPLATE = '台水資訊整合系統_設計文件_{date}.docx'",
        "",
        "# 每日排程時間（Daemon 模式）",
        "SCHEDULE_TIME = '10:00'",
        "",
        "# 保留最近 N 份文件（0 = 永遠保留）",
        "KEEP_LAST_N = 30",
    ], caption="doc_generator.py 設定常數")

    b.spacer(4)
    b.h2("6.2 新增要跳過的資料夾")
    b.body("若有新的備份資料夾需要排除，修改 analyzer.py 的兩個常數：")
    b.code([
        "# 完全跳過的資料夾名稱（精確比對）",
        "SKIP_DIRS = {",
        "    '.git', '.claude', 'target', 'build', '__pycache__',",
        "    'node_modules', '.tmp_ccms', 'update', 'archive',",
        "    'myNewBackupFolder',   # <-- 在此加入新的排除資料夾",
        "}",
        "",
        "# 資料夾名稱含有以下字串時跳過（不分大小寫）",
        "COPY_FOLDER_MARKERS = ('ciopy', ' - copy', '-copy', '_copy', '複製', 'backup', 'bak')",
    ], caption="analyzer.py 過濾設定")

    b.spacer(4)
    b.h2("6.3 修改文件主題色彩")
    b.body("所有色彩定義在 word_builder.py 頂部的 Palette 區段：")
    b.code([
        "C_TITLE      = RGBColor(0x1F, 0x4E, 0x79)   # 封面標題色（深藍）",
        "C_H1         = RGBColor(0x1F, 0x4E, 0x79)   # H1 標題色",
        "C_H2         = RGBColor(0x2E, 0x75, 0xB6)   # H2 標題色（中藍）",
        "C_HDR_FILL   = 'D5E8F0'                      # 表格標題列背景色",
        "C_ALT_FILL   = 'EBF3FB'                      # 表格交錯列背景色",
    ], caption="word_builder.py 色彩設定")

    b.spacer(4)
    b.h2("6.4 新增靜態章節內容")
    b.body("以「已知限制」章節為例，在 word_builder.py 的對應位置找到 rows 陣列並新增一行：")
    b.code([
        "# word_builder.py 第 10 章「已知限制」",
        "b.table(",
        "    ['項目', '說明'],",
        "    [",
        "        ['API 端點停用',    '...'],",
        "        ['無 Token Revoke', '...'],",
        "        ['新增限制項目',    '在此加入新的限制說明'],  # <-- 新增這行",
        "    ],",
        "    [4.5, 11.5]",
        ")",
    ], caption="新增靜態內容範例")

    _page_break(b.doc)

    # ══ 7. 執行紀錄 ════════════════════════════════════════════════════════════
    b.h1("7. 執行紀錄（Log）")

    b.h2("7.1 Log 位置")
    b.code([
        "docs/generated/doc_gen.log",
        "",
        "# 範例內容：",
        "2026-05-18 10:00:01 [INFO] Analysing project: D:\\...",
        "2026-05-18 10:00:01 [INFO] Found 3 modules, 28 classes, 3 JSP files.",
        "2026-05-18 10:00:02 [INFO] Building Word document → ...\\20260518.docx",
        "2026-05-18 10:00:06 [INFO] Document saved: ...\\20260518.docx  (50 KB)",
    ])

    b.h2("7.2 查看最近 30 行 Log")
    b.code([
        "# PowerShell：",
        "Get-Content 'docs\\generated\\doc_gen.log' -Tail 30",
        "",
        "# cmd.exe：",
        "type docs\\generated\\doc_gen.log",
    ])

    b.h2("7.3 Windows Task Scheduler 執行結果碼")
    b.table(
        ["LastTaskResult 代碼", "說明", "處理方式"],
        [
            ["0",       "成功",                              "正常"],
            ["267011",  "尚未執行過（初始狀態）",             "第一次排程後的正常狀態"],
            ["1",       "Python 腳本發生 Exception",          "查看 doc_gen.log 確認錯誤原因"],
            ["2147942402", "找不到 Python 或腳本路徑",        "確認 Python 在 PATH 中，或重新執行 setup_scheduler.ps1"],
        ],
        [4.0, 5.5, 6.5]
    )

    _page_break(b.doc)

    # ══ 8. 常見問題 ════════════════════════════════════════════════════════════
    b.h1("8. 常見問題（FAQ）")

    faq = [
        (
            "Q: 執行時顯示 ModuleNotFoundError: No module named 'docx'",
            "尚未安裝 python-docx。",
            ["pip install python-docx"]
        ),
        (
            "Q: PermissionError — 文件無法寫入",
            "目標 .docx 檔案正在被 Word 開啟。",
            [
                "# 解決方法一：關閉 Word 中的文件後重新執行",
                "# 解決方法二：程式會自動加上時間戳記產生新檔名，不影響正常執行",
            ]
        ),
        (
            "Q: 排程執行成功但文件沒有更新（LastTaskResult = 0，但看不到新檔案）",
            "確認輸出目錄是否正確。",
            [
                "# 查看排程執行的工作目錄與指令",
                "Get-ScheduledTask 'TaiwanWater_DocGen' | Select -ExpandProperty Actions",
                "",
                "# 確認輸出目錄",
                "dir docs\\generated\\",
            ]
        ),
        (
            "Q: 文件中偵測到 0 個模組",
            "專案根目錄設定不正確，或所有來源目錄都被過濾掉了。",
            [
                "# 確認掃描路徑",
                "python doc_generator.py --once --root \"D:\\正確的專案路徑\"",
                "",
                "# 快速診斷：直接執行 analyzer.py",
                "python doc_gen\\analyzer.py \"D:\\project\\Ai3\\台水\\water\"",
            ]
        ),
        (
            "Q: 想在 Linux / macOS 上排程",
            "改用 cron 代替 Windows Task Scheduler。",
            [
                "# 編輯 crontab",
                "crontab -e",
                "",
                "# 加入以下行（每天早上 10:00 執行）",
                "0 10 * * * cd /path/to/doc_gen && python doc_generator.py --once >> /tmp/doc_gen.log 2>&1",
            ]
        ),
        (
            "Q: 如何確認掃描到正確的檔案？",
            "直接執行 analyzer.py 查看詳細掃描結果。",
            [
                "python doc_gen\\analyzer.py \"D:\\project\\Ai3\\台水\\water\"",
                "",
                "# 輸出範例：",
                "# Modules: ['ecpsso', 'waterCusSSO', 'waterHrSync']",
                "# [ecpsso]       unit_id=df6c8bac-...  classes=10  jsps=0",
                "# [waterCusSSO]  unit_id=              classes=0   jsps=3",
                "# [waterHrSync]  unit_id=4806684c-...  classes=18  jsps=0",
                "# DB tables found: ['TcChatMan', 'TcChatMessage', 'TcContact', ...]",
            ]
        ),
    ]

    for q, reason, code_lines in faq:
        b.h3(q)
        b.body(f"原因：{reason}")
        b.code(code_lines)
        b.spacer(2)

    _page_break(b.doc)

    # ══ 9. 技術架構圖 ══════════════════════════════════════════════════════════
    b.h1("9. doc-gen 自身架構圖")

    b.h2("9.1 資料流")
    b.code([
        "                    doc_generator.py（主程式）",
        "                            │",
        "              ┌─────────────┴──────────────┐",
        "              ▼                            ▼",
        "       analyzer.py                  word_builder.py",
        "    （靜態程式碼分析）             （Word 文件建構）",
        "              │                            │",
        "    ┌─────────┼──────────┐                 │",
        "    ▼         ▼          ▼                 ▼",
        " .java 檔   .jsp 檔  .properties      python-docx",
        " （regex)  （regex)   （key=value)    （XML 操作）",
        "    │         │          │                 │",
        "    └────┬────┘          │                 │",
        "         ▼               │                 │",
        "    ProjectInfo ◀────────┘                 │",
        "    （資料模型）                              │",
        "         │                                 │",
        "         └──────────────────────────────── ▶",
        "                                     台水資訊整合系統_設計文件_YYYYMMDD.docx",
    ])

    b.spacer(4)
    b.h2("9.2 排程觸發流程")
    b.code([
        "每日 10:00",
        "    │",
        "    ▼",
        "Windows Task Scheduler",
        "  任務名稱: TaiwanWater_DocGen",
        "  指令: python doc_generator.py --once",
        "    │",
        "    ▼",
        "doc_generator.generate()",
        "  Step 1: analyzer.analyze()     ← 掃描原始碼（0.1 秒）",
        "  Step 2: word_builder.build()   ← 產生 Word（4-5 秒）",
        "  Step 3: _rotate_old_docs()     ← 刪除 30 天前舊文件",
        "  Step 4: 寫入 doc_gen.log",
        "    │",
        "    ▼",
        "docs/generated/台水資訊整合系統_設計文件_YYYYMMDD.docx  ✓",
    ])

    b.spacer(4)
    b.h2("9.3 執行時間分析")
    b.table(
        ["執行階段", "時間", "說明"],
        [
            ["目錄遍歷（rglob）",        "~0.01 秒", "找到所有 .java / .jsp / .properties 檔案"],
            ["正規表達式解析（28 檔）",   "~0.09 秒", "提取類別、方法、SQL、API、Unit ID 等資訊"],
            ["Word 文件建構",             "~4-5 秒",  "python-docx 建立表格、標題、頁首頁尾、ASCII 圖"],
            ["檔案寫入 / 重新命名",       "~0.1 秒",  "原子性寫檔（先寫 .tmp.docx，再重新命名）"],
            ["合計",                     "< 6 秒",   "28 個 Java 檔 + 3 個 JSP 檔的專案規模"],
        ],
        [5.0, 2.0, 9.0]
    )
    b.note("執行時間主要瓶頸在 Word 文件建構（python-docx），而非程式碼掃描。"
           "即使專案成長到 500 個 Java 檔，總執行時間預估仍在 30 秒以內。")

    # ══ Footer note ════════════════════════════════════════════════════════════
    b.spacer(16)
    hr = b.doc.add_paragraph(); _para_border_bottom(hr, GRY_COLOR, 4)
    p = b.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"doc-gen User Guide v1.0 — 產生時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    r.font.size=Pt(9); r.font.name=FONT; r.font.color.rgb=C_GRAY; r.font.italic=True

    b.save(OUTPUT)


if __name__ == "__main__":
    build()
