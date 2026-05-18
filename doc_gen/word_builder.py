"""
word_builder.py
---------------
Builds a professional Word document from ProjectInfo metadata.
Uses python-docx. Install: pip install python-docx
"""

from __future__ import annotations
import os
from pathlib import Path
from datetime import datetime
from typing import List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml.ns as ns

from analyzer import ProjectInfo, Module, JavaClass, JspEndpoint


# ── Colour palette ─────────────────────────────────────────────────────────────
C_TITLE      = RGBColor(0x1F, 0x4E, 0x79)   # dark blue  - cover title
C_H1         = RGBColor(0x1F, 0x4E, 0x79)   # dark blue  - heading 1
C_H2         = RGBColor(0x2E, 0x75, 0xB6)   # mid  blue  - heading 2
C_H3         = RGBColor(0x1F, 0x4E, 0x79)   # dark blue  - heading 3
C_HDR_FILL   = "D5E8F0"                      # light blue - table header bg
C_ALT_FILL   = "EBF3FB"                      # pale blue  - alt row bg
C_ACCENT     = "2E75B6"                      # mid  blue  - border accent
C_GRAY       = "808080"                      # gray       - footer text
C_CODE_BG    = "F2F2F2"                      # light gray - code block bg
C_NOTE       = RGBColor(0xC5, 0x5A, 0x11)   # orange     - warning note

FONT_BODY    = "微軟正黑體"
FONT_CODE    = "Courier New"
FONT_COVER   = "微軟正黑體"


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "BFBFBF"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_para_border_bottom(para, color: str = "2E75B6", size: int = 12):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_col_width(table, col_idx: int, width_cm: float):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)


def _add_page_break(doc: Document):
    para = doc.add_paragraph()
    run  = para.add_run()
    br   = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def _add_header_footer(doc: Document, title: str):
    section = doc.sections[0]

    # Header
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run(title)
    run.font.size   = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.name   = FONT_BODY
    _set_para_border_bottom(hp, color="BFBFBF", size=4)

    # Footer with page number
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Page X of Y" using field codes
    for text, field_name in [
        ("第 ", None), (None, "PAGE"), (" / ", None), (None, "NUMPAGES"), (" 頁", None)
    ]:
        if text:
            run = fp.add_run(text)
        else:
            run = fp.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)

            instrText = OxmlElement('w:instrText')
            instrText.text = f' {field_name} '
            run._r.append(instrText)

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run2 = fp.add_run()
            run2._r.append(fldChar2)

        run.font.size  = Pt(9)
        run.font.name  = FONT_BODY
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ── Paragraph / Table factories ────────────────────────────────────────────────

class DocBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        doc = self.doc
        # Page: A4, margins 2.5 cm
        section = doc.sections[0]
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)

        # Normal style base
        normal = doc.styles['Normal']
        normal.font.name = FONT_BODY
        normal.font.size = Pt(10.5)

    # ── Text helpers ────────────────────────────────────────────────────────────

    def cover_title(self, text: str, size: int = 28):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name  = FONT_COVER
        r.font.size  = Pt(size)
        r.font.bold  = True
        r.font.color.rgb = C_TITLE
        return p

    def cover_sub(self, text: str, size: int = 14):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name  = FONT_COVER
        r.font.size  = Pt(size)
        r.font.color.rgb = C_H2
        return p

    def meta_line(self, label: str, value: str):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}：")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.name = FONT_BODY
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.name = FONT_BODY
        pf = p.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after  = Pt(4)

    def h1(self, text: str):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(18)
        pf.space_after  = Pt(6)
        r = p.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(16)
        r.font.bold  = True
        r.font.color.rgb = C_H1
        _set_para_border_bottom(p, color=C_ACCENT, size=8)
        # mark as Heading 1 for TOC
        p.style = self.doc.styles['Heading 1']
        p.clear()
        pf = p.paragraph_format
        pf.space_before = Pt(18)
        pf.space_after  = Pt(6)
        r = p.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(16)
        r.font.bold  = True
        r.font.color.rgb = C_H1
        _set_para_border_bottom(p, color=C_ACCENT, size=8)
        return p

    def h2(self, text: str):
        p = self.doc.add_paragraph(style='Heading 2')
        p.clear()
        pf = p.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after  = Pt(4)
        r = p.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(13)
        r.font.bold  = True
        r.font.color.rgb = C_H2
        return p

    def h3(self, text: str):
        p = self.doc.add_paragraph(style='Heading 3')
        p.clear()
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.name  = FONT_BODY
        r.font.size  = Pt(11)
        r.font.bold  = True
        r.font.color.rgb = C_H3
        return p

    def body(self, text: str, bold: bool = False, italic: bool = False):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.name   = FONT_BODY
        r.font.size   = Pt(10.5)
        r.font.bold   = bold
        r.font.italic = italic
        return p

    def bullet(self, text: str):
        p = self.doc.add_paragraph(style='List Bullet')
        r = p.add_run(text)
        r.font.name = FONT_BODY
        r.font.size = Pt(10.5)
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after  = Pt(1)
        return p

    def numbered(self, text: str):
        p = self.doc.add_paragraph(style='List Number')
        r = p.add_run(text)
        r.font.name = FONT_BODY
        r.font.size = Pt(10.5)
        return p

    def code_block(self, lines: list[str]):
        for line in lines:
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after  = Pt(0)
            pf.left_indent  = Cm(1)
            r = p.add_run(line)
            r.font.name = FONT_CODE
            r.font.size = Pt(9)
            # shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  C_CODE_BG)
            pPr.append(shd)

    def note(self, text: str):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent  = Cm(0.5)
        pf.space_before = Pt(4)
        pf.space_after  = Pt(4)
        r1 = p.add_run("⚠ 注意：")
        r1.font.bold = True
        r1.font.color.rgb = C_NOTE
        r1.font.name = FONT_BODY
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.color.rgb = C_NOTE
        r2.font.name = FONT_BODY
        r2.font.size = Pt(10)

    def spacer(self, pts: float = 6):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(pts)
        p.paragraph_format.space_after  = Pt(0)

    # ── Table helper ────────────────────────────────────────────────────────────

    def table(self, headers: List[str], rows: List[List[str]],
              col_widths_cm: List[float]):
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # set col widths
        for ci, w in enumerate(col_widths_cm):
            _set_col_width(tbl, ci, w)

        # header row
        hdr_row = tbl.rows[0]
        for ci, hdr in enumerate(headers):
            cell = hdr_row.cells[ci]
            cell.text = ''
            _set_cell_bg(cell, C_HDR_FILL)
            _set_cell_borders(cell, C_ACCENT)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            r = p.add_run(hdr)
            r.font.bold  = True
            r.font.name  = FONT_BODY
            r.font.size  = Pt(10)
            r.font.color.rgb = C_H1

        # data rows
        for ri, row_data in enumerate(rows):
            row = tbl.add_row()
            alt  = (ri % 2 == 1)
            for ci, val in enumerate(row_data):
                cell = row.cells[ci]
                cell.text = ''
                if alt:
                    _set_cell_bg(cell, C_ALT_FILL)
                _set_cell_borders(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                r = p.add_run(str(val) if val else '—')
                r.font.name = FONT_BODY
                r.font.size = Pt(10)

        return tbl

    def ascii_diagram(self, lines: List[str], title: str = ""):
        if title:
            self.body(title, bold=True)
        self.code_block(lines)

    def save(self, path: str):
        self.doc.save(path)
        return path


# ── ASCII Diagrams ─────────────────────────────────────────────────────────────

def build_system_context(info: ProjectInfo) -> List[str]:
    mod_names = [m.name for m in info.modules]
    return [
        "┌─────────────────────────────────────────────────────────────────────┐",
        "│              台灣自來水公司 資訊整合系統  ( System Context )            │",
        "├──────────────────────────────────────────────────────────────────────┤",
        "│  使用者            系統邊界                         外部系統           │",
        "│                                                                     │",
        "│  ┌─────────┐  HTTPS  ┌─────────────────┐  LDAP  ┌───────────────┐  │",
        "│  │ 內部員工  │────────▶│  waterCusSSO     │───────▶│ Active Dir.   │  │",
        "│  └─────────┘         │  (WAR / JSP)     │        └───────────────┘  │",
        "│                      └────────┬─────────┘                           │",
        "│                     In-Proc   │                                     │",
        "│                      ┌────────▼─────────┐  JDBC  ┌───────────────┐  │",
        "│  ┌─────────┐         │    ecpsso.jar     │───────▶│   ECP 主 DB   │  │",
        "│  │ 客服人員  │─Web──▶ │  (SSO 核心模組)   │        └───────────────┘  │",
        "│  └─────────┘         └──────────────────┘                           │",
        "│                                                                     │",
        "│                      ┌──────────────────┐  JDBC  ┌───────────────┐  │",
        "│                      │  waterHrSync.jar  │───────▶│  KM 人事 DB   │  │",
        "│                      │  (HR 同步 + 聊天)  │        └───────────────┘  │",
        "│                      │  每日 02:00 排程   │  JDBC  ┌───────────────┐  │",
        "│                      │                   │───────▶│  客服聊天 DB   │  │",
        "│                      └──────────────────┘        └───────────────┘  │",
        "└─────────────────────────────────────────────────────────────────────┘",
    ]


def build_deployment_diagram() -> List[str]:
    return [
        "┌─────────────────── 台水內部網路 ─────────────────────────────────────┐",
        "│                                                                      │",
        "│  員工 Browser ──HTTPS:443──▶ JBoss / WildFly 應用伺服器              │",
        "│                              waterCusSSO.war                        │",
        "│                              ├─ ecpsso.jar                          │",
        "│                              └─ waterHrSync.jar                     │",
        "│                                    │                                │",
        "│          ┌─────────────────────────┼──────────────────────┐         │",
        "│          ▼                          ▼                      ▼         │",
        "│  ┌──────────────┐      ┌─────────────────┐    ┌──────────────────┐  │",
        "│  │  ECP 主 DB   │      │   KM 人事 DB     │    │  Active Directory │  │",
        "│  │ executor:    │      │ executor:km      │    │  (LDAP:636)      │  │",
        "│  │ default(讀寫)│      │ (唯讀)            │    └──────────────────┘  │",
        "│  └──────────────┘      └─────────────────┘                          │",
        "│  ┌──────────────┐                                                    │",
        "│  │  客服聊天 DB  │ ◀── JDBC 唯讀                                      │",
        "│  └──────────────┘                                                    │",
        "└──────────────────────────────────────────────────────────────────────┘",
    ]


def build_sso_flow() -> List[str]:
    return [
        "  員工Browser       login.jsp      AD/LDAP       sso.jsp     SSOService    ECP-DB",
        "      │                 │              │              │            │            │",
        "  ①帳密登入            │              │              │            │            │",
        "      │──POST{loginName,password}──▶  │              │            │            │",
        "      │                 │──LDAP bind──▶│              │            │            │",
        "      │                 │◀─ 驗證結果 ──│              │            │            │",
        "      │                 │─────────────────────────────────────────────────────▶│",
        "      │◀─ {success,token}              │              │            │            │",
        "      │                 │              │              │            │            │",
        "  ②後續SSO驗證          │              │              │            │            │",
        "      │────────────────────────────────────── POST{loginName,token} ──────────▶│",
        "      │                 │              │              │─doSSOToken─▶            │",
        "      │                 │              │              │            │─查詢5分鐘─▶│",
        "      │                 │              │              │            │◀── token ──│",
        "      │   Token有效 ◀──────────────────────────────── success      │            │",
        "      │   Token失效 ── AD Fallback ────▶              fail         │            │",
        "      │                 │              │              │            │            │",
        "  ③稽核日誌             │              │              │─addLog()───▶            │",
        "      │                 │              │              │            │──INSERT──▶ │",
    ]


def build_hr_sync_flow() -> List[str]:
    return [
        "  Timer(02:00)   HrSyncService    HrSyncDao      KM 人事DB     ECP Database",
        "       │               │               │               │              │",
        "       │─syncHrData()─▶│               │               │              │",
        "       │               │──getHrUserList()──▶           │              │",
        "       │               │               │──SELECT v_KM_USER(00*)──▶    │",
        "       │               │               │◀──────── HR員工清單 ──────────│",
        "       │               │──getHrDeptList()──▶           │              │",
        "       │               │               │──SELECT v_KM_DEPT──▶         │",
        "       │               │◀──部門清單 ────│◀──────────────────────────── │",
        "       │               │  拓撲排序(BFS)  │               │              │",
        "       │               │─[loop 每個部門]─────────────────────────────▶ │",
        "       │               │               │               │  upsert TsDept│",
        "       │               │─[loop 每個員工]─────────────────────────────▶ │",
        "       │               │               │               │  INSERT/UPDATE │",
        "       │               │               │               │  TsUser+Acct  │",
        "       │               │               │               │  assignRoles  │",
        "       │               │               │               │  (IT→SysAdmin)│",
        "       │◀─ 完成(新增N/更新M) ───────────│               │              │",
    ]


def build_er_diagram(db_tables: List[str]) -> List[str]:
    lines = [
        "  ┌────────────────┐     ┌────────────────────┐     ┌───────────────┐",
        "  │   TsUser       │     │   TsDepartment      │     │   TsAccount   │",
        "  │────────────────│     │────────────────────│     │───────────────│",
        "  │ FId (PK)       │     │ FId (PK)            │     │ FId (PK)      │",
        "  │ FName          │     │ FName               │     │ FLoginName    │",
        "  │ FDepartmentId──│─────▶ FParentId (self)    │     │ FPassword     │",
        "  │ FEnabled       │     │ FTreeLevel          │     │ FEmail        │",
        "  └───────┬────────┘     └────────────────────┘     └──────┬────────┘",
        "          │                                                 │",
        "          └────────────▶ TsAccountIdentity ◀───────────────┘",
        "                         (FAccountId, FEntityId)",
        "",
        "  ┌────────────────┐     ┌─────────────┐",
        "  │   TsRole       │     │ TsRoleUser  │",
        "  │────────────────│     │─────────────│",
        "  │ FId (PK)       │◀────│ FRoleId(FK) │",
        "  │ FName          │     │ FUserId(FK) │",
        "  └────────────────┘     └─────────────┘",
        "",
        "  ┌──────────────────────────────────┐   ┌────────────────────────────────┐",
        "  │   TpCUSmLoginLog (自訂)           │   │ TpCUSmChatMessageSummary (自訂) │",
        "  │──────────────────────────────────│   │────────────────────────────────│",
        "  │ FId, U_SessionID, U_token        │   │ FId, U_AgentId, U_ChatId       │",
        "  │ FCreateUserId, FCreateTime ◀索引  │   │ U_Content, U_SendTime          │",
        "  └──────────────────────────────────┘   │ U_UserId, U_UserName           │",
        "                                          └────────────────────────────────┘",
    ]
    if db_tables:
        lines += ["", f"  實際偵測到的資料表（共 {len(db_tables)} 個）："]
        line = "  "
        for t in db_tables:
            if len(line) + len(t) + 2 > 80:
                lines.append(line)
                line = "  "
            line += t + "  "
        if line.strip():
            lines.append(line)
    return lines


# ── Main Document Generator ────────────────────────────────────────────────────

def build_document(info: ProjectInfo, output_path: str) -> str:
    b = DocBuilder()
    doc = b.doc

    _add_header_footer(doc, "台灣自來水公司 資訊整合系統 — 架構與設計文件")

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    b.spacer(60)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)

    b.cover_title("台灣自來水公司", 28)
    b.cover_title("資訊整合系統", 24)
    b.spacer(6)
    b.cover_sub("系統架構與詳細設計文件", 16)
    b.spacer(4)

    # horizontal rule
    hr = doc.add_paragraph()
    _set_para_border_bottom(hr, C_ACCENT, size=12)

    b.spacer(8)
    b.meta_line("版本", "Auto-generated v2.0")
    b.meta_line("產生時間", info.generated_at)
    b.meta_line("分析路徑", info.root)
    b.meta_line("狀態", "自動產生文件")
    b.spacer(4)

    _add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 1 – ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    b.h1("第一部分：系統架構圖")

    b.h2("1. 系統整體架構（System Context）")
    b.body("呈現台水資訊整合系統與所有外部系統的邊界關係及主要使用者。")
    b.spacer(2)
    b.ascii_diagram(build_system_context(info))

    b.h2("2. 容器 / 模組架構")
    b.body(f"本系統共偵測到 {len(info.modules)} 個模組，部署於同一 JBoss/WildFly 應用伺服器。")
    rows = []
    for m in info.modules:
        cls_cnt = len(m.classes)
        jsp_cnt = len(m.jsps)
        rows.append([m.name, m.description,
                     f"Java: {cls_cnt}, JSP: {jsp_cnt}",
                     m.unit_id or '—'])
    b.table(
        ["模組名稱", "說明", "檔案數", "Unit ID"],
        rows,
        [3.5, 5.5, 2.5, 5.0]
    )

    b.h2("3. SSO 登入流程圖（Sequence）")
    b.ascii_diagram(build_sso_flow())

    b.h2("4. HR 同步流程圖（Sequence）")
    b.ascii_diagram(build_hr_sync_flow())

    b.h2("5. 資料庫 ER 圖")
    b.ascii_diagram(build_er_diagram(info.db_tables))

    b.h2("6. 部署架構圖")
    b.ascii_diagram(build_deployment_diagram())

    _add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PART 2 – DETAILED DESIGN
    # ══════════════════════════════════════════════════════════════════════════
    b.h1("第二部分：詳細設計文件")

    # ── Section 1: 專案概述 ───────────────────────────────────────────────────
    b.h2("1. 專案概述")
    b.h3("1.1 系統目標")
    for goal in [
        "單一登入（SSO）：員工只需一組帳號即可存取 ECP 平台，透過 Token 機制串接外部系統。",
        "HR 資料同步：每日自動將 KM 人事系統的員工與部門資料同步至 ECP，保持帳號資料一致性。",
        "聊天摘要彙整：每日彙整客服聊天記錄，儲存至 ECP 資料庫供後續分析。",
    ]:
        b.bullet(goal)

    b.h3("1.2 系統邊界")
    b.table(
        ["系統", "角色", "擁有方"],
        [
            ["ECP (Quicksilver 8.5)", "主要企業協作平台，本系統的部署容器", "Chainsea 鏈鎖科技"],
            ["KM 人事系統", "HR 資料來源（唯讀）", "台水 IT"],
            ["Active Directory", "身分驗證（LDAP）", "台水 IT"],
            ["客服聊天系統", "聊天資料來源（唯讀）", "台水 IT"],
            ["本系統", "整合中介層", "Ai3 / 台水"],
        ],
        [5.0, 6.5, 4.5]
    )

    # ── Section 2: 技術堆疊 ───────────────────────────────────────────────────
    b.h2("2. 技術堆疊")
    b.table(
        ["層級", "技術", "版本"],
        [
            ["語言", "Java", "17"],
            ["建置工具", "Maven", "3.x"],
            ["應用框架", "Quicksilver + ECP", "7.1.31.beta22 + 8.5.02.36"],
            ["應用伺服器", "JBoss / WildFly", "（依 ECP 平台版本）"],
            ["前端", "JSP (Jakarta EE)", "4.0"],
            ["Servlet", "Jakarta Servlet", "4.0"],
            ["JSON 處理", "org.json", "20240303"],
            ["日誌", "SLF4J + Log4j2", "2.0.16"],
            ["資料庫", "MariaDB / SQL Server", "（依部署環境）"],
            ["身分驗證", "LDAP / Active Directory", "—"],
            ["部署單元", "WAR + JAR (plugin)", "—"],
        ],
        [4.0, 5.5, 6.5]
    )

    # ── Section 3: 模組說明 ───────────────────────────────────────────────────
    b.h2("3. 模組說明（自動分析結果）")

    for mod in info.modules:
        if mod.name == 'other':
            continue
        b.h3(f"3.{info.modules.index(mod)+1}  {mod.name} — {mod.description}")
        b.body(f"Unit ID：{mod.unit_id or '（未偵測到）'}")

        # Classes by layer
        layer_map: dict[str, list[JavaClass]] = {}
        for jc in mod.classes:
            layer_map.setdefault(jc.layer, []).append(jc)

        if layer_map:
            layer_order = ['home', 'model', 'dao', 'service', 'api', 'action', 'timer', 'other']
            rows = []
            for lay in layer_order:
                for jc in layer_map.get(lay, []):
                    methods_str = ', '.join(jc.methods[:5])
                    if len(jc.methods) > 5:
                        methods_str += f' … (+{len(jc.methods)-5})'
                    rows.append([
                        lay.upper(),
                        f"{jc.kind} {jc.name}",
                        methods_str or '—',
                    ])
            if rows:
                b.table(["層次", "類別 / 介面", "主要方法"], rows, [2.0, 4.5, 9.5])

        # API endpoints
        all_apis = [(jc.name, path) for jc in mod.classes for path in jc.api_paths]
        if all_apis:
            b.body("偵測到的 @Api 端點：", bold=True)
            b.table(
                ["類別", "API Path"],
                all_apis,
                [5.0, 11.0]
            )

        # Timer schedules
        timers = [(jc.name, jc.timer_schedule) for jc in mod.classes if jc.timer_schedule]
        if timers:
            b.body("排程設定：", bold=True)
            b.table(["Timer Class", "Cron 表達式"], timers, [7.0, 9.0])

        # JSP endpoints
        if mod.jsps:
            b.body(f"JSP 端點（共 {len(mod.jsps)} 個）：", bold=True)
            rows = []
            for je in mod.jsps:
                rows.append([
                    je.name + '.jsp',
                    je.http_method,
                    '是' if je.cors else '否',
                    ', '.join(je.request_fields[:5]) or '—',
                    ', '.join(je.flow_steps[:3]) or '—',
                ])
            b.table(
                ["檔案", "HTTP", "CORS", "Request 欄位", "呼叫服務"],
                rows,
                [3.0, 1.5, 1.5, 4.0, 6.0]
            )

        # Properties
        for pf in mod.properties:
            if pf.entries:
                b.body(f"設定檔：{Path(pf.filepath).name}", bold=True)
                entries = [[k, v if 'password' not in k.lower() else '（已加密）']
                           for k, v in pf.entries.items()]
                b.table(["屬性", "值"], entries, [8.0, 8.0])

        b.spacer(4)

    # ── Section 4: 資料庫設計 ─────────────────────────────────────────────────
    b.h2("4. 資料庫設計")

    b.h3("4.1 資料庫連線設定")
    b.table(
        ["Executor", "用途", "資料庫"],
        [
            ["default", "ECP 主資料庫（讀寫）", "ECP DB"],
            ["km",      "KM 人事系統（唯讀）",  "KM DB"],
            ["（預設）", "客服聊天資料庫（唯讀）", "Chat DB"],
        ],
        [3.5, 5.5, 7.0]
    )

    b.h3("4.2 自訂資料表")
    b.body("TpCUSmLoginLog — SSO 登入稽核", bold=True)
    b.table(
        ["欄位", "型別", "說明"],
        [
            ["FId",                "UUID PK",    "主鍵"],
            ["U_SessionID",        "VARCHAR",    "ECP Session ID"],
            ["U_token",            "VARCHAR",    "SSO Token"],
            ["FName",              "VARCHAR",    "登入狀態（LoginStatus）"],
            ["FCreateUserId",      "UUID FK",    "登入使用者 ID（→ TsUser）"],
            ["FCreateDepartmentId","UUID FK",    "所屬部門 ID（→ TsDepartment）"],
            ["U_CreateUserName",   "VARCHAR",    "使用者名稱（冗餘備份）"],
            ["U_CreateDepartment", "VARCHAR",    "部門名稱（冗餘備份）"],
            ["FCreateTime",        "DATETIME",   "建立時間 ★ 需建索引（5 分鐘查詢）"],
        ],
        [4.5, 2.5, 9.0]
    )
    b.note("FCreateTime 需建立索引，getSSOToken() 的 5 分鐘範圍查詢依賴此欄位效能。")
    b.spacer(4)

    b.body("TpCUSmChatMessageSummary — 聊天摘要", bold=True)
    b.table(
        ["欄位", "型別", "說明"],
        [
            ["FId",         "UUID PK",   "主鍵"],
            ["U_AgentId",   "VARCHAR",   "客服坐席 ID"],
            ["U_AgentName", "VARCHAR",   "客服坐席姓名"],
            ["U_ChatId",    "VARCHAR",   "聊天室 ID"],
            ["U_Content",   "TEXT",      "訊息內容"],
            ["U_RoomId",    "VARCHAR",   "聊天房間 ID"],
            ["U_SenderId",  "VARCHAR",   "發送者 ID"],
            ["U_SenderName","VARCHAR",   "發送者姓名"],
            ["U_SendTime",  "DATETIME",  "發送時間"],
            ["U_Type",      "VARCHAR",   "訊息類型"],
            ["U_UserId",    "VARCHAR",   "客戶使用者 ID"],
            ["U_UserName",  "VARCHAR",   "客戶使用者姓名"],
        ],
        [4.5, 2.5, 9.0]
    )

    b.h3("4.3 ECP 平台核心資料表（參考）")
    b.table(
        ["資料表", "說明"],
        [
            ["TsUser",            "ECP 使用者主表"],
            ["TsDepartment",      "部門樹狀結構"],
            ["TsAccount",         "登入帳號"],
            ["TsAccountIdentity", "帳號 ↔ 人員對應（支援多身分）"],
            ["TsRole",            "角色定義"],
            ["TsRoleUser",        "角色指派關係"],
        ],
        [5.0, 11.0]
    )

    if info.db_tables:
        b.h3("4.4 自動偵測到的資料表")
        b.body(f"以下 {len(info.db_tables)} 個資料表從 DAO 的 SQL 語句中自動偵測：")
        rows = [[t] for t in sorted(info.db_tables)]
        b.table(["資料表名稱"], rows, [16.0])

    # ── Section 5: API 規格 ───────────────────────────────────────────────────
    b.h2("5. API 規格")

    b.h3("5.1 Active Endpoints（JSP）")
    b.table(
        ["端點", "HTTP", "說明", "Request", "Response"],
        [
            ["/custom/login/login.jsp", "POST", "帳密登入（CORS）",
             '{"loginName":"…","password":"…"}', '{"success":true}'],
            ["/custom/login/sso.jsp",   "POST", "SSO Token 驗證",
             '{"loginName":"…","token":"…"}',    '{"success":true/false}'],
            ["/custom/login/autologin.jsp","GET","自動登入輔助", "—", "—"],
        ],
        [4.5, 1.5, 3.5, 4.5, 2.0]
    )

    b.h3("5.2 Disabled Endpoints（暫未啟用）")
    b.table(
        ["Path", "方法", "說明"],
        [
            ["/TCBEcpSSO/EcpSSO",   "POST", "doSSO() ECP API 版本（@Api，isTokenRequired=false）"],
            ["/TCBEcpSSO/SSOToken", "POST", "doSSOToken() ECP API 版本"],
        ],
        [4.5, 1.5, 10.0]
    )
    b.note("這些 Endpoint 程式碼已存在但被停用，功能目前由 JSP 直接承擔。")

    # ── Section 6: 排程設計 ───────────────────────────────────────────────────
    b.h2("6. 排程設計")
    b.table(
        ["排程名稱", "Class", "執行時間", "說明"],
        [
            ["HR 同步",  "SyncHrDataRountine",           "每日 02:00", "同步人員 / 部門至 ECP"],
            ["聊天摘要", "GetChatMessageSummaryRountine", "每日（預設昨日）", "彙整前一日聊天記錄"],
        ],
        [3.0, 5.5, 2.5, 5.0]
    )
    b.spacer(2)
    for note_txt in [
        "HR 同步排定在 04:00 備份作業之前執行，確保備份時資料已是最新狀態。",
        "兩個排程不應同時執行，避免資料庫鎖衝突。",
        "GetChatMessageSummaryRountine 支援傳入指定日期參數（格式：yyyy-MM-dd）。",
    ]:
        b.bullet(note_txt)

    # ── Section 7: 安全設計 ───────────────────────────────────────────────────
    b.h2("7. 安全設計")

    b.h3("7.1 Token 安全機制")
    b.table(
        ["機制", "說明"],
        [
            ["有效期限",  "Token 5 分鐘後自動失效"],
            ["綁定使用者", "Token 與 empId 綁定，無法跨使用者使用"],
            ["Fallback",  "Token 失效時嘗試 AD 重新驗證，而非直接放行"],
            ["稽核日誌",  "所有驗證嘗試（成功 / 失敗）均寫入 TpCUSmLoginLog 與 sso_auth.log"],
        ],
        [3.5, 12.5]
    )

    b.h3("7.2 密碼安全")
    b.table(
        ["項目", "說明"],
        [
            ["預設密碼", "儲存於 hrsync.properties，以加密格式儲存（ECP 平台加密機制）"],
            ["AD 驗證",  "密碼不落地，由 LDAP bind 操作完成驗證"],
            ["LDAP 連線", "建議使用 LDAPS（port 636）加密傳輸"],
        ],
        [3.5, 12.5]
    )

    b.h3("7.3 資料隔離")
    b.bullet("KM 人事資料庫使用獨立 Executor（km），僅有讀取權限。")
    b.bullet("客服聊天資料庫同樣為唯讀存取。")
    b.bullet("CORS：login.jsp 需確認 Access-Control-Allow-Origin 鎖定特定網域而非開放 *。")

    # ── Section 8: 錯誤處理 ───────────────────────────────────────────────────
    b.h2("8. 錯誤處理與日誌")

    b.h3("8.1 日誌策略")
    b.table(
        ["日誌", "位置", "內容"],
        [
            ["sso_auth.log",   "應用伺服器日誌目錄", "SSO Token 驗證結果、來源 IP、時間戳"],
            ["SLF4J / Log4j2", "ECP 標準日誌",       "HR 同步執行結果（新增 N 筆、更新 M 筆）"],
            ["JSP stdout",     "伺服器 stdout",      "登入請求的時間戳、IP、使用者名稱"],
        ],
        [3.5, 4.0, 8.5]
    )

    b.h3("8.2 HR 同步錯誤處理")
    b.table(
        ["情境", "處理方式"],
        [
            ["KM DB 連線失敗",    "拋出例外，本次同步中止，記錄 ERROR log"],
            ["部門父節點不存在",   "拓撲排序保證父節點先建立，若仍失敗則跳過並記錄警告"],
            ["使用者資料衝突",     "以 loginName 為唯一鍵，EXISTS 判斷後做 INSERT 或 UPDATE"],
            ["角色指派失敗",       "記錄警告，不中止整體同步流程"],
        ],
        [4.5, 11.5]
    )

    b.h3("8.3 聊天摘要錯誤處理")
    b.table(
        ["情境", "處理方式"],
        [
            ["無聊天記錄",  "返回空集合，不執行 INSERT，記錄 INFO log"],
            ["資料重複插入", "以 FId UUID 為主鍵，重跑時若 FId 重複會拋 PK 違反例外"],
        ],
        [4.5, 11.5]
    )
    b.note("聊天摘要模組目前缺乏冪等設計，建議加入 DELETE + INSERT 或 UPSERT 機制。")

    # ── Section 9: 部署 ────────────────────────────────────────────────────────
    b.h2("9. 部署與設定")

    b.h3("9.1 部署步驟")
    for step in [
        "編譯 ecpsso      → 產生 ecpsso.jar",
        "編譯 waterHrSync → 產生 waterHrSync.jar",
        "將 JAR 複製至 waterCusSSO/WEB-INF/lib/",
        "打包 waterCusSSO → 產生 waterCusSSO.war",
        "部署 WAR 至 JBoss/WildFly",
        "在 ECP 後台設定 Timer Routine（HR sync 02:00、Chat summary 每日）",
    ]:
        b.numbered(step)

    b.h3("9.2 hrsync.properties")
    b.code_block([
        "# TsAccountIdentity 使用的身分類型 UUID（請勿任意修改）",
        "accountIdentityTypeId=564cf69e-76d6-4baf-b584-6e04c2911dae",
        "",
        "# 新建帳號的預設密碼（ECP 加密格式）",
        "accountDefaultPassword=<encrypted>",
        "",
        "# ECP 部門樹根節點 UUID",
        "TsDepartmentRootFId=00000000-0000-0000-1001-000000000001",
    ])

    b.h3("9.3 環境需求")
    b.table(
        ["項目", "需求"],
        [
            ["Java",          "17+"],
            ["JBoss/WildFly", "ECP 8.5 相容版本"],
            ["ECP",           "8.5.02.36"],
            ["Quicksilver",   "7.1.31.beta22"],
            ["網路",           "應用伺服器可連線至 AD、KM DB、Chat DB"],
        ],
        [4.0, 12.0]
    )

    # ── Section 10: 已知限制 ──────────────────────────────────────────────────
    b.h2("10. 已知限制與改善建議")

    b.h3("10.1 已知限制")
    b.table(
        ["項目", "說明"],
        [
            ["API 端點停用",   "SSOLoginLogApi 的 ECP API 端點目前為空實作，功能由 JSP 承擔"],
            ["無 Token Revoke", "Token 無法主動撤銷，只能等待 5 分鐘自然過期"],
            ["聊天摘要非冪等", "重複執行會產生重複資料"],
            ["CORS 白名單",    "login.jsp 的 CORS 設定需確認是否鎖定特定來源域名"],
            ["帳號停用邏輯",   "KM 人員離職後，ECP 帳號的停用/刪除流程未在程式碼中看到"],
        ],
        [4.5, 11.5]
    )

    b.h3("10.2 建議改善項目")
    b.table(
        ["優先級", "項目", "說明"],
        [
            ["高", "聊天摘要冪等設計", "執行前先 DELETE 當日資料，或改用 UPSERT"],
            ["高", "帳號離職停用",     "增加「HR 中不存在的 ECP 帳號自動停用」邏輯"],
            ["中", "Token 延長機制",  "提供 refresh token 或延長有效期至 30 分鐘"],
            ["中", "CORS 白名單",     "限制 Access-Control-Allow-Origin 為特定域名"],
            ["低", "API 端點啟用",    "啟用 ECP API 版本的 SSO 端點，取代 JSP 直接處理"],
            ["低", "單元測試",        "補充 HrSyncService、SSOLoginLogService 的單元測試"],
        ],
        [1.5, 4.0, 10.5]
    )

    # footer note
    b.spacer(12)
    hr2 = doc.add_paragraph()
    _set_para_border_bottom(hr2, color="BFBFBF", size=4)
    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_note.add_run(
        f"本文件由自動化程式依據原始碼分析產生（{info.generated_at}），如有疑義請以原始碼為準。"
    )
    r.font.size = Pt(9)
    r.font.name = FONT_BODY
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    r.font.italic = True

    b.save(output_path)
    return output_path


# ── CLI test ───────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    import sys
    from analyzer import ProjectAnalyzer
    root = sys.argv[1] if len(sys.argv) > 1 else '.'
    out  = sys.argv[2] if len(sys.argv) > 2 else 'output.docx'
    info = ProjectAnalyzer(root).analyze()
    build_document(info, out)
    print(f"Saved: {out}")
