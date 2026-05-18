"""
report_builder.py
-----------------
Builds a Word (.docx) TOGAF compliance report from ProjectCheckResult.

Colour scheme
  🔴 HIGH     → red fill
  🟡 MEDIUM   → yellow fill
  🟢 LOW      → light-blue fill  (informational)
  ✅ PASS     → green fill

Usage:
    from report_builder import build_report
    build_report(project_check_result, "output.docx")
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from checker import ProjectCheckResult, PrincipleResult, Finding
from principles_config import get_recommendations


# ── Colour constants (RGB) ────────────────────────────────────────────────────

C_RED    = RGBColor(0xC0, 0x00, 0x00)   # dark red text
C_AMBER  = RGBColor(0xFF, 0x8C, 0x00)   # dark orange text
C_GREEN  = RGBColor(0x37, 0x86, 0x26)   # dark green text
C_BLUE   = RGBColor(0x1F, 0x49, 0x7D)   # dark blue text
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK  = RGBColor(0x00, 0x00, 0x00)

FILL_RED    = 'C00000'   # cell background
FILL_AMBER  = 'FFC000'
FILL_YELLOW = 'FFFF00'
FILL_GREEN  = '70AD47'
FILL_BLUE   = 'BDD7EE'
FILL_LGRAY  = 'F2F2F2'
FILL_DGRAY  = '595959'
FILL_HEADER = '1F497D'   # navy

SEVERITY_FILL = {
    'HIGH':   FILL_RED,
    'MEDIUM': FILL_AMBER,
    'LOW':    FILL_BLUE,
}
SEVERITY_LABEL = {
    'HIGH':   '高風險',
    'MEDIUM': '中風險',
    'LOW':    '低風險',
}

PRIORITY_FILL = {
    '最高': FILL_RED,
    '高':   FILL_AMBER,
}


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_fill(cell, hex_color: str):
    """Set table cell background fill (shading)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """kwargs: top/bottom/left/right — each = {'val':'single','sz':'4','color':'auto'}"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, attrs in kwargs.items():
        el = OxmlElement(f'w:{side}')
        for k, v in attrs.items():
            el.set(qn(f'w:{k}'), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_align(para, alignment):
    para.alignment = alignment


def _run_font(run, size_pt: int, bold=False, color: RGBColor = None, italic=False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _cell_para(cell, text: str, size_pt=10, bold=False,
               color: RGBColor = C_BLACK, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.space_before = Pt(2)
    para.space_after  = Pt(2)
    run = para.add_run(text)
    _run_font(run, size_pt, bold=bold, color=color)


def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    h.space_before = Pt(12 if level == 1 else 6)
    h.space_after  = Pt(4)
    return h


def _add_para(doc: Document, text: str, size_pt=10, color=C_BLACK, bold=False):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after  = Pt(2)
    run = p.add_run(text)
    _run_font(run, size_pt, bold=bold, color=color)
    return p


def _page_break(doc: Document):
    doc.add_page_break()


# ── Table builder helpers ─────────────────────────────────────────────────────

def _make_table(doc: Document, rows: int, cols: int,
                col_widths_cm: List[float] = None) -> object:
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths_cm:
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[j])
    return tbl


def _header_row(tbl, col_idx: int, headers: List[str],
                fill=FILL_HEADER, text_color=C_WHITE, size_pt=10):
    row = tbl.rows[col_idx]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _set_cell_fill(cell, fill)
        _cell_para(cell, h, size_pt=size_pt, bold=True, color=text_color,
                   align=WD_ALIGN_PARAGRAPH.CENTER)


# ── Score bar (ASCII-style via table) ────────────────────────────────────────

def _score_bar_text(score: int) -> str:
    filled = round(score / 5)       # out of 20 blocks
    empty  = 20 - filled
    return f"{'█' * filled}{'░' * empty}  {score}%"


def _score_fill(score: int) -> str:
    if score >= 80:
        return FILL_GREEN
    if score >= 60:
        return FILL_AMBER
    return FILL_RED


def _score_text_color(score: int) -> RGBColor:
    if score >= 80:
        return C_GREEN
    if score >= 60:
        return C_AMBER
    return C_RED


# ── Header / Footer ──────────────────────────────────────────────────────────

def _add_header_footer(doc: Document, title: str):
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(title)
    _run_font(run, 9, italic=True, color=RGBColor(0x70, 0x70, 0x70))

    # Footer: page numbers via field
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = fp.add_run('第 ')
    _run_font(run, 9, color=RGBColor(0x70, 0x70, 0x70))

    # PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = fp.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    _run_font(run2, 9, color=RGBColor(0x70, 0x70, 0x70))

    run3 = fp.add_run(' 頁，共 ')
    _run_font(run3, 9, color=RGBColor(0x70, 0x70, 0x70))

    # NUMPAGES field
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = ' NUMPAGES '
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run4 = fp.add_run()
    run4._r.append(fldChar3)
    run4._r.append(instrText2)
    run4._r.append(fldChar4)
    _run_font(run4, 9, color=RGBColor(0x70, 0x70, 0x70))

    run5 = fp.add_run(' 頁')
    _run_font(run5, 9, color=RGBColor(0x70, 0x70, 0x70))


# ── Cover page ───────────────────────────────────────────────────────────────

def _build_cover(doc: Document, pr: ProjectCheckResult):
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('台灣自來水公司')
    _run_font(run, 28, bold=True, color=C_BLUE)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run('TOGAF 軟體開發原則合規報告')
    _run_font(run, 22, bold=True, color=C_BLUE)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'生成時間：{pr.generated_at}')
    _run_font(run, 12, color=RGBColor(0x40, 0x40, 0x40))

    root_para = doc.add_paragraph()
    root_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = root_para.add_run(f'掃描目錄：{pr.root}')
    _run_font(run, 11, color=RGBColor(0x40, 0x40, 0x40))

    files_para = doc.add_paragraph()
    files_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = files_para.add_run(f'掃描檔案總數：{pr.total_files_scanned} 個')
    _run_font(run, 11, color=RGBColor(0x40, 0x40, 0x40))

    doc.add_paragraph()
    doc.add_paragraph()

    # Overall score gauge
    scores = [r.score for r in pr.results]
    overall = round(sum(scores) / len(scores)) if scores else 0

    gauge_para = doc.add_paragraph()
    gauge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = gauge_para.add_run(f'整體合規分數  {overall}%')
    _run_font(run, 20, bold=True, color=_score_text_color(overall))

    bar_para = doc.add_paragraph()
    bar_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = bar_para.add_run(_score_bar_text(overall))
    _run_font(run, 14, color=_score_text_color(overall))

    _page_break(doc)


# ── Executive summary ─────────────────────────────────────────────────────────

def _build_summary(doc: Document, pr: ProjectCheckResult):
    _add_heading(doc, '1. 執行摘要', 1)

    _add_para(doc,
              f'本報告針對「台灣自來水公司 資訊整合系統」原始碼進行靜態分析，'
              f'共掃描 {pr.total_files_scanned} 個檔案，'
              f'依 TOGAF 架構框架定義之 8 項軟體開發準則逐一評估。',
              size_pt=11)
    doc.add_paragraph()

    # Summary table
    tbl = _make_table(doc, rows=len(pr.results) + 1, cols=6,
                      col_widths_cm=[1.5, 2.0, 3.5, 2.5, 2.2, 2.0])
    _header_row(tbl, 0, ['優先度', '準則ID', '準則名稱', '掃描行數', '發現問題數', '合規分數'])

    for i, r in enumerate(pr.results, start=1):
        row = tbl.rows[i]
        prio_fill = PRIORITY_FILL.get(r.priority, FILL_LGRAY)

        _set_cell_fill(row.cells[0], prio_fill)
        _cell_para(row.cells[0], r.priority, bold=True, color=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

        _cell_para(row.cells[1], r.principle_id, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[2], r.principle_name)

        _cell_para(row.cells[3], str(r.checks_run),
                   align=WD_ALIGN_PARAGRAPH.CENTER)

        high_cnt = sum(1 for f in r.findings if f.severity == 'HIGH')
        med_cnt  = sum(1 for f in r.findings if f.severity == 'MEDIUM')
        low_cnt  = sum(1 for f in r.findings if f.severity == 'LOW')
        finding_text = (
            f'高:{high_cnt}  中:{med_cnt}  低:{low_cnt}'
            if r.findings else '無'
        )
        fill = FILL_RED if high_cnt else (FILL_AMBER if med_cnt else (FILL_LGRAY if not r.findings else FILL_BLUE))
        _set_cell_fill(row.cells[4], fill)
        _cell_para(row.cells[4], finding_text, align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=C_WHITE if (high_cnt or med_cnt) else C_BLACK)

        score_fill = _score_fill(r.score)
        _set_cell_fill(row.cells[5], score_fill)
        _cell_para(row.cells[5], f'{r.score}%', bold=True,
                   color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # Legend
    legend = doc.add_paragraph()
    legend.add_run('說明：').bold = True
    legend.add_run(
        ' 優先度「最高(🔴)」為 TOGAF 核心安全與互通性準則；「高(🟠)」為架構品質準則。'
        '合規分數 ≥80% 為綠色（合格），60-79% 為橙色（需改善），<60% 為紅色（不合格）。'
        '掃描行數為工具實際逐行比對的次數，僅供參考；'
        '合規分數採嚴重度扣分制：以「規則＋檔案」為單位，同一問題在同一檔案多行只扣一次，'
        '出現在不同檔案則各扣一次（HIGH -20 / MEDIUM -10 / LOW -4），與掃描行數無關。'
    )
    _page_break(doc)


# ── Principle detail section ─────────────────────────────────────────────────

def _build_principle_section(doc: Document, r: PrincipleResult, section_no: int,
                              project_root: str = ""):
    status = 'PASS ✓' if r.passed else 'FAIL ✗'
    status_color = C_GREEN if r.passed else C_RED

    _add_heading(doc, f'{section_no}. {r.principle_id} — {r.principle_name}', 1)

    # Metadata row
    meta_tbl = _make_table(doc, rows=1, cols=4,
                            col_widths_cm=[3.0, 3.0, 4.5, 5.5])
    row = meta_tbl.rows[0]
    _set_cell_fill(row.cells[0], PRIORITY_FILL.get(r.priority, FILL_LGRAY))
    _cell_para(row.cells[0], f'優先度：{r.priority}', bold=True, color=C_WHITE,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    status_fill = FILL_GREEN if r.passed else FILL_RED
    _set_cell_fill(row.cells[1], status_fill)
    _cell_para(row.cells[1], status, bold=True, color=C_WHITE,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    _cell_para(row.cells[2], f'掃描行數：{r.checks_run}',
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_fill(row.cells[3], _score_fill(r.score))
    _cell_para(row.cells[3], _score_bar_text(r.score), bold=True, color=C_WHITE,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    if not r.findings:
        p = doc.add_paragraph()
        run = p.add_run('  本準則未發現任何問題，符合 TOGAF 要求。')
        _run_font(run, 11, color=C_GREEN, bold=True)
        doc.add_paragraph()
        return

    # Findings table
    _add_para(doc, '發現問題清單：', size_pt=11, bold=True)

    tbl = _make_table(doc, rows=len(r.findings) + 1, cols=5,
                      col_widths_cm=[1.8, 5.0, 1.2, 5.5, 3.0])
    _header_row(tbl, 0, ['嚴重等級', '檔案路徑', '行號', '問題說明', '違規代碼片段'])

    for i, f in enumerate(r.findings, start=1):
        row = tbl.rows[i]
        sev_fill = SEVERITY_FILL.get(f.severity, FILL_BLUE)

        _set_cell_fill(row.cells[0], sev_fill)
        label = SEVERITY_LABEL.get(f.severity, f.severity)
        _cell_para(row.cells[0], label, bold=True, color=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

        # Show path relative to project root so module name is always visible.
        # Fall back to the raw filepath only for sentinel "(全域)" entries.
        try:
            display_fp = str(Path(f.filepath).relative_to(project_root)) \
                         if project_root and f.filepath != '(全域)' \
                         else f.filepath
        except ValueError:
            display_fp = f.filepath   # not under root — show as-is
        _cell_para(row.cells[1], display_fp, size_pt=9)

        line_text = str(f.line_no) if f.line_no > 0 else '-'
        _cell_para(row.cells[2], line_text, align=WD_ALIGN_PARAGRAPH.CENTER)

        _cell_para(row.cells[3], f.rule, size_pt=10)

        # Evidence: monospace-ish, truncate
        ev = (f.evidence[:100] + '…') if len(f.evidence) > 100 else f.evidence
        ev_cell = row.cells[4]
        ev_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        ev_para = ev_cell.paragraphs[0]
        ev_para.space_before = Pt(1)
        ev_para.space_after  = Pt(1)
        ev_run = ev_para.add_run(ev)
        ev_run.font.size = Pt(8)
        ev_run.font.name = 'Courier New'

        # Alternate row shading
        if i % 2 == 0:
            _set_cell_fill(row.cells[1], FILL_LGRAY)
            _set_cell_fill(row.cells[2], FILL_LGRAY)
            _set_cell_fill(row.cells[3], FILL_LGRAY)
            _set_cell_fill(row.cells[4], FILL_LGRAY)

    doc.add_paragraph()

    # Recommendation
    _add_para(doc, '改善建議：', size_pt=11, bold=True)
    recs = get_recommendations(r.principle_id)
    for rec in recs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(rec)
        _run_font(run, 10)

    doc.add_paragraph()


# Recommendations are now managed in principles_config.py — imported above.


# ── Appendix: statistics ──────────────────────────────────────────────────────

def _build_appendix(doc: Document, pr: ProjectCheckResult):
    _add_heading(doc, f'{len(pr.results) + 2}. 附錄：統計摘要', 1)

    # Severity breakdown table
    _add_para(doc, '問題嚴重等級分佈：', size_pt=11, bold=True)
    tbl = _make_table(doc, rows=len(pr.results) + 1, cols=5,
                      col_widths_cm=[4.5, 2.5, 2.5, 2.5, 3.0])
    _header_row(tbl, 0, ['準則', '高風險', '中風險', '低風險', '合計'])
    for i, r in enumerate(pr.results, start=1):
        row = tbl.rows[i]
        high = sum(1 for f in r.findings if f.severity == 'HIGH')
        med  = sum(1 for f in r.findings if f.severity == 'MEDIUM')
        low  = sum(1 for f in r.findings if f.severity == 'LOW')
        total = len(r.findings)
        _cell_para(row.cells[0], f'{r.principle_id} {r.principle_name}')
        _set_cell_fill(row.cells[1], FILL_RED if high else FILL_LGRAY)
        _cell_para(row.cells[1], str(high), bold=high > 0, color=C_WHITE if high else C_BLACK,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_fill(row.cells[2], FILL_AMBER if med else FILL_LGRAY)
        _cell_para(row.cells[2], str(med), bold=med > 0, color=C_WHITE if med else C_BLACK,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[3], str(low), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[4], str(total), align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    _add_para(doc,
              f'本報告由 TOGAF Checker 自動生成 · 生成時間：{pr.generated_at} · '
              f'工具版本：1.0.0',
              size_pt=9, color=RGBColor(0x70, 0x70, 0x70))


# ── Main entry ────────────────────────────────────────────────────────────────

def build_report(pr: ProjectCheckResult, out_path: str):
    """Build the Word report and save to *out_path*."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    report_title = f'台水資訊整合系統 TOGAF 合規報告  ·  {pr.generated_at}'
    _add_header_footer(doc, report_title)

    # Cover
    _build_cover(doc, pr)

    # Executive summary
    _build_summary(doc, pr)

    # One section per principle
    for idx, r in enumerate(pr.results, start=2):
        _build_principle_section(doc, r, idx, project_root=pr.root)
        _page_break(doc)

    # Appendix
    _build_appendix(doc, pr)

    doc.save(out_path)
