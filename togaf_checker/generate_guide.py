"""
generate_guide.py
-----------------
Standalone script that generates the TOGAF Checker User Guide in Word format.
No AI is used at runtime — pure Python + python-docx.

Usage:
    python generate_guide.py
    python generate_guide.py --out D:/output
"""

from __future__ import annotations

import argparse
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x1F, 0x49, 0x7D)
C_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
C_GREEN  = RGBColor(0x37, 0x86, 0x26)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
C_AMBER  = RGBColor(0xFF, 0x8C, 0x00)
C_GRAY   = RGBColor(0x40, 0x40, 0x40)
C_LGRAY  = RGBColor(0x70, 0x70, 0x70)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_BLACK  = RGBColor(0x00, 0x00, 0x00)

F_NAVY   = '1F497D'
F_BLUE   = '2E74B5'
F_GREEN  = '70AD47'
F_AMBER  = 'FFC000'
F_RED    = 'C00000'
F_LGRAY  = 'F2F2F2'
F_DGRAY  = 'D9D9D9'
F_TEAL   = '00B0F0'
F_WHITE  = 'FFFFFF'

TODAY = datetime.now().strftime('%Y-%m-%d %H:%M')

# ── XML helpers ────────────────────────────────────────────────────────────────

def _shd(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _cell(cell, text: str, size=10, bold=False, color=C_BLACK,
          align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment  = align
    p.space_before = Pt(2)
    p.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _tbl(doc, rows, cols, widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if widths:
        for row in t.rows:
            for j, c in enumerate(row.cells):
                if j < len(widths):
                    c.width = Cm(widths[j])
    return t


def _hdr(tbl, col_labels, fill=F_NAVY):
    row = tbl.rows[0]
    for i, lbl in enumerate(col_labels):
        _shd(row.cells[i], fill)
        _cell(row.cells[i], lbl, bold=True, color=C_WHITE,
              align=WD_ALIGN_PARAGRAPH.CENTER)


def _h(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.space_before = Pt(14 if level == 1 else 8 if level == 2 else 4)
    h.space_after  = Pt(4)
    return h


def _p(doc, text='', size=11, color=C_BLACK, bold=False, italic=False,
       align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.alignment    = align
    p.space_before = Pt(2)
    p.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        run.font.color.rgb = color
    return p


def _bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.space_before = Pt(1)
    p.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _numbered(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Number')
    p.space_before = Pt(1)
    p.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def _code(doc, text):
    """Mono-spaced code block (single paragraph, Courier New)."""
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after  = Pt(2)
    # Light gray shading via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p


def _page_break(doc):
    doc.add_page_break()


def _add_header_footer(doc, title):
    for section in doc.sections:
        # Header
        h = section.header
        h.is_linked_to_previous = False
        hp = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(title)
        run.font.size   = Pt(9)
        run.font.italic = True
        run.font.color.rgb = C_LGRAY

        # Footer with page numbers
        f = section.footer
        f.is_linked_to_previous = False
        fp = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _fld(p, instruction):
            run = p.add_run()
            run.font.size = Pt(9)
            run.font.color.rgb = C_LGRAY
            r = run._r
            b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); r.append(b)
            i = OxmlElement('w:instrText'); i.text = instruction; r.append(i)
            e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end');   r.append(e)

        for txt in ['第 ', 'PAGE', ' 頁，共 ', 'NUMPAGES', ' 頁']:
            if txt in ('PAGE', 'NUMPAGES'):
                _fld(fp, f' {txt} ')
            else:
                run = fp.add_run(txt)
                run.font.size = Pt(9)
                run.font.color.rgb = C_LGRAY


# ══════════════════════════════════════════════════════════════════════════════
# Chapter builders
# ══════════════════════════════════════════════════════════════════════════════

# ── Cover ─────────────────────────────────────────────────────────────────────

def _cover(doc):
    for _ in range(4):
        _p(doc)

    t = _p(doc, '台灣自來水公司', size=26, bold=True, color=C_NAVY,
           align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, 'TOGAF 原則合規檢查工具', size=22, bold=True, color=C_BLUE,
       align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, 'User Guide  使用者指南', size=16, color=C_GRAY,
       align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(2):
        _p(doc)

    _p(doc, f'文件版本：1.0.2        生成時間：{TODAY}',
       size=11, color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, '工具位置：<專案根目錄>/togaf_checker/',
       size=11, color=C_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(3):
        _p(doc)

    # Tagline box (simulate with a 1-cell table)
    tb = _tbl(doc, 1, 1, [16])
    _shd(tb.rows[0].cells[0], F_NAVY)
    _cell(tb.rows[0].cells[0],
          '100% offline  ·  No AI at runtime  ·  Pure Python regex static analysis',
          size=12, bold=True, color=C_WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    _page_break(doc)


# ── Chapter 1: Introduction ───────────────────────────────────────────────────

def _ch1_intro(doc):
    _h(doc, '1. 工具簡介', 1)

    _p(doc,
       'TOGAF Checker 是一套純 Python 靜態程式碼分析工具，'
       '用於評估「台灣自來水公司 資訊整合系統」的原始碼是否符合 '
       'TOGAF（The Open Group Architecture Framework）所定義的 8 項軟體開發準則。'
       '工具每天自動執行一次，掃描 .java、.jsp、.properties 檔案，'
       '並以 Word 格式輸出彩色合規報告。',
       size=11)

    _p(doc)

    # Key facts box
    tb = _tbl(doc, 5, 2, [5, 11.5])
    facts = [
        ('執行環境',  'Python 3.9+  ·  Windows 10/11'),
        ('輸入',      '.java / .jsp / .properties 原始碼'),
        ('輸出',      'docs/generated/台水TOGAF合規報告_YYYYMMDD_HHMM.docx'),
        ('排程',      'Windows Task Scheduler，每天 10:00 自動執行'),
        ('AI 使用',   '❌ 無 — 完全離線，不呼叫任何 AI API'),
    ]
    for i, (k, v) in enumerate(facts):
        _shd(tb.rows[i].cells[0], F_LGRAY)
        _cell(tb.rows[i].cells[0], k, bold=True, size=10)
        _cell(tb.rows[i].cells[1], v, size=10)

    _p(doc)


# ── Chapter 2: Does it use Claude Code / AI? ─────────────────────────────────

def _ch2_ai(doc):
    _h(doc, '2. 這個工具有用到 AI 或 Claude Code 嗎？', 1)

    _h(doc, '2.1  執行期間：完全不使用 AI', 2)
    _p(doc,
       '工具在每天 10:00 自動掃描並產出報告的整個過程中，'
       '完全不呼叫任何 AI 服務、不連線至網路、不使用 Claude API。'
       '所有分析邏輯都是由固定的正規表達式（regex）規則引擎完成。',
       size=11)

    _p(doc)

    # Comparison table
    _h(doc, '2.2  AI vs 本工具分析方式對照', 2)
    tb = _tbl(doc, 5, 3, [4.5, 5.5, 6.5])
    _hdr(tb, ['比較項目', 'AI 程式碼審查', 'TOGAF Checker（本工具）'])
    rows = [
        ('分析方法',   '大型語言模型推理（LLM inference）',
                       '固定 regex 模式匹配 + 計數統計'),
        ('是否聯網',   '是（呼叫 Claude / GPT API）',
                       '否（完全離線）'),
        ('重複性',     '相同程式碼每次可能得到不同措詞',
                       '完全確定性，相同輸入永遠相同輸出'),
        ('規則來源',   '模型訓練資料（不透明）',
                       '人工定義的 TOGAF 準則 → 可審計'),
    ]
    for i, (a, b, c) in enumerate(rows, start=1):
        if i % 2 == 0:
            _shd(tb.rows[i].cells[0], F_LGRAY)
            _shd(tb.rows[i].cells[1], F_LGRAY)
            _shd(tb.rows[i].cells[2], F_LGRAY)
        _cell(tb.rows[i].cells[0], a, bold=True, size=10)
        _cell(tb.rows[i].cells[1], b, size=10)
        _cell(tb.rows[i].cells[2], c, size=10)

    _p(doc)

    _h(doc, '2.3  Claude Code 的角色', 2)
    _p(doc,
       'Claude Code 是 Anthropic 提供的 AI 輔助開發工具（CLI）。'
       '它被用來「撰寫」本工具的原始碼，就像一位資深工程師協助設計並實作整個專案。'
       '但一旦程式碼寫好並交付後，Claude Code 便不再涉入——',
       size=11)

    tb2 = _tbl(doc, 3, 2, [3, 13.5])
    phases = [
        ('開發期（過去）',
         'Claude Code 根據需求設計架構、撰寫 checker.py / report_builder.py / main.py，'
         '並在 git worktree 中測試修正。'),
        ('交付後（現在）',
         '工具獨立運行。每天 10:00 由 Windows Task Scheduler 呼叫 python main.py，'
         '完全不需要 Claude Code 存在。'),
        ('結論',
         'Claude Code = 建造工具的工匠；本工具 = 建好後獨立運作的機器。'
         '二者在執行期間毫無關係。'),
    ]
    fills = [F_BLUE, F_GREEN, F_DGRAY]
    for i, (phase, desc) in enumerate(phases):
        _shd(tb2.rows[i].cells[0], fills[i])
        _cell(tb2.rows[i].cells[0], phase, bold=True, color=C_WHITE, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(tb2.rows[i].cells[1], desc, size=10)

    _p(doc)
    _page_break(doc)


# ── Chapter 3: Tech Stack ─────────────────────────────────────────────────────

def _ch3_stack(doc):
    _h(doc, '3. 技術棧（Tech Stack）', 1)

    _h(doc, '3.1  Runtime 依賴', 2)
    tb = _tbl(doc, 8, 4, [3.5, 2.5, 4.5, 6])
    _hdr(tb, ['元件', '版本', '用途', '說明'])
    rows = [
        ('Python',         '3.9+',   'Runtime',
         '標準直譯器；stdlib 已包含 re、pathlib、dataclasses、argparse'),
        ('python-docx',    '≥1.1.0', 'Word 文件產生',
         '以 Python API 操作 .docx XML；唯一的第三方依賴'),
        ('re（stdlib）',   '內建',   '靜態程式碼分析',
         '所有 TOGAF 規則均以正規表達式實作，無需額外 NLP 函式庫'),
        ('pathlib（stdlib）','內建', '檔案系統遞迴掃描',
         'Path.rglob() 遍歷 .java/.jsp/.properties'),
        ('dataclasses（stdlib）','內建','資料模型',
         'Finding / PrincipleResult / ProjectCheckResult 結構定義'),
        ('argparse（stdlib）','內建','CLI 介面',
         '--root / --out / --time / --daemon 參數'),
        ('Windows Task Scheduler','OS 內建','排程執行',
         '由 setup_scheduler.ps1 以 PowerShell 註冊，每天 10:00 觸發'),
    ]
    for i, r in enumerate(rows, start=1):
        if i % 2 == 0:
            for c in tb.rows[i].cells:
                _shd(c, F_LGRAY)
        _cell(tb.rows[i].cells[0], r[0], bold=True, size=10)
        _cell(tb.rows[i].cells[1], r[1], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(tb.rows[i].cells[2], r[2], size=10)
        _cell(tb.rows[i].cells[3], r[3], size=10)

    _p(doc)

    _h(doc, '3.2  開發期工具（執行期不需要）', 2)
    tb2 = _tbl(doc, 4, 3, [4, 3, 9.5])
    _hdr(tb2, ['工具', '類別', '用途'])
    dev_tools = [
        ('Claude Code (Anthropic)', 'AI 輔助開發 CLI', '撰寫並測試本工具的所有原始碼'),
        ('Git',                     '版本控制',        '原始碼管理；工具在 .claude/worktrees/ 下開發'),
        ('PowerShell 5.1+',         'Windows Shell',   'setup_scheduler.ps1 註冊排程工作'),
    ]
    for i, (a, b, c) in enumerate(dev_tools, start=1):
        _cell(tb2.rows[i].cells[0], a, bold=True, size=10)
        _cell(tb2.rows[i].cells[1], b, size=10)
        _cell(tb2.rows[i].cells[2], c, size=10)

    _p(doc)

    _h(doc, '3.3  TOGAF 準則對應的規則技術', 2)
    _p(doc,
       '每一條 TOGAF 準則都對應一組 Python regex 模式，'
       '以下列出各準則的核心偵測技術：',
       size=11)

    tb3 = _tbl(doc, 9, 3, [2, 4, 10.5])
    _hdr(tb3, ['準則', '準則名稱', '偵測技術'])
    rules = [
        ('P15', 'Data Security',
         'HARDCODED_PW: password\\s*=\\s*"[^"]{3,}"  ·  SQL_CONCAT: SELECT.*\\+\\s*\\w+  ·  '
         'TOKEN_FALSE: isTokenRequired=false  ·  CORS_WILDCARD: Access-Control-Allow-Origin.*\\*'),
        ('P16', 'Technology Independence',
         'DIRECT_IMPL: new\\s+\\w*(?:Impl|DAO)\\(  ·  VENDOR_LOCK: import com.mysql/oracle.jdbc  ·  '
         'interface-to-class 比例計算  ·  @Autowired/@Inject 存在性偵測'),
        ('P20', 'Control Technical Diversity',
         '6 種 JSON 函式庫 import 模式  ·  4 種 Log 框架 import 模式  ·  '
         '5 種 HTTP 用戶端 import 模式；計算各類別共存數量'),
        ('P21', 'Interoperability',
         'HARDCODED_IP: http://\\d{1,3}\\.\\d{1,3}\\...  ·  '
         'JSP 輸出 JSON 但無 Content-Type header  ·  @Api Swagger 標記存在性'),
        ('P6',  'Service Orientation',
         '4 層分層 class/interface 模式偵測（Model/DAO/Service/Action）  ·  '
         'method count > 25 → 過胖類別警告'),
        ('P5',  'Common Use Applications',
         '跨檔案 class 名稱 Map 建立  ·  檢測相同類別名稱存在於不同模組路徑  ·  '
         'common/util/shared 套件存在性'),
        ('P4',  'Business Continuity',
         'EMPTY_CATCH: catch\\s*\\([^)]+\\)\\s*\\{[\\s]*\\}  ·  '
         'BROAD_CATCH: catch\\(Exception e\\) 後 5 行無 log  ·  null 防禦計數'),
        ('P7',  'Compliance with Law',
         'PII_LOG: log.*password/idno/身分證  ·  PERSONAL_DATA_LOG: log.*name/address  ·  '
         'AUDIT_LOG: audit.*( 存在性  ·  MASK_PATTERN: mask/脫敏/replaceAll.*\\*'),
    ]
    fills_alt = [F_WHITE, F_LGRAY]
    for i, (pid, pname, tech) in enumerate(rules, start=1):
        fill = fills_alt[i % 2]
        _shd(tb3.rows[i].cells[0], F_NAVY if i % 2 else F_BLUE)
        _cell(tb3.rows[i].cells[0], pid, bold=True, color=C_WHITE,
              size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        if fill != F_WHITE:
            _shd(tb3.rows[i].cells[1], fill)
            _shd(tb3.rows[i].cells[2], fill)
        _cell(tb3.rows[i].cells[1], pname, bold=True, size=10)
        # code-like text for the patterns
        p = tb3.rows[i].cells[2].paragraphs[0]
        p.space_before = Pt(2)
        p.space_after  = Pt(2)
        run = p.add_run(tech)
        run.font.size = Pt(8.5)
        run.font.name = 'Courier New'

    _p(doc)
    _page_break(doc)


# ── Chapter 4: How to Run ─────────────────────────────────────────────────────

def _ch4_howto(doc):
    _h(doc, '4. 執行方法', 1)

    _h(doc, '4.1  前置需求', 2)
    _bullet(doc, 'Python 3.9 或以上（確認：python --version）')
    _bullet(doc, '安裝 python-docx：pip install python-docx')
    _bullet(doc, 'Windows 10/11（排程功能）；Linux/Mac 可改用 cron')

    _p(doc)

    _h(doc, '4.2  目錄結構', 2)
    _code(doc,
          '<專案根目錄>/\n'
          '├── togaf_checker/\n'
          '│   ├── principles_config.py  # ★ 原則定義檔（唯一修改入口）\n'
          '│   ├── checker.py            # 8 項 TOGAF 準則規則引擎\n'
          '│   ├── report_builder.py     # Word 報告產生器\n'
          '│   ├── main.py               # 入口程式（CLI + 排程）\n'
          '│   ├── generate_guide.py     # 本使用者指南產生腳本\n'
          '│   ├── setup_scheduler.bat   # 排程設定啟動器（雙擊）\n'
          '│   ├── setup_scheduler.ps1   # 排程設定主程式（PowerShell）\n'
          '│   └── requirements.txt      # python-docx>=1.1.0\n'
          '└── docs/\n'
          '    └── generated/\n'
          '        └── 台水TOGAF合規報告_YYYYMMDD_HHMM.docx  # 輸出報告'
    )
    _p(doc)

    _h(doc, '4.3  手動執行（一次性掃描）', 2)
    _p(doc, '開啟命令提示字元（cmd）或 PowerShell，進入 togaf_checker 目錄：', size=11)
    _code(doc, 'cd D:\\your_project\\togaf_checker')
    _p(doc, '使用預設設定（掃描上一層目錄，報告存至 docs/generated/）：', size=11)
    _code(doc, 'python main.py')
    _p(doc, '指定掃描目標與輸出位置：', size=11)
    _code(doc, 'python main.py --root D:\\project\\water --out D:\\reports')
    _p(doc)

    _h(doc, '4.4  排程設定（每天 10:00 自動執行）', 2)

    _p(doc, '方法 A：雙擊批次檔（推薦）', size=11, bold=True)
    _numbered(doc, '在 Windows 檔案總管中找到 togaf_checker\\setup_scheduler.bat')
    _numbered(doc, '雙擊執行，PowerShell 視窗會開啟並自動完成設定')
    _numbered(doc, '視窗最後會詢問「是否立即執行測試？」，輸入 Y 可驗證')
    _p(doc)

    _p(doc, '方法 B：直接執行 PowerShell 腳本', size=11, bold=True)
    _code(doc, 'powershell -ExecutionPolicy Bypass -File setup_scheduler.ps1')
    _p(doc)

    _p(doc, '方法 C：Daemon 模式（程式持續常駐）', size=11, bold=True)
    _code(doc, 'python main.py --daemon\n'
               'python main.py --daemon --time 09:30    # 改為 09:30 執行')
    _p(doc)

    # Task Scheduler details table
    _h(doc, '4.5  Windows Task Scheduler 工作詳情', 2)
    tb = _tbl(doc, 6, 2, [5, 11.5])
    details = [
        ('工作名稱',  'TaiwanWater_TOGAFChecker'),
        ('觸發程序',  '每天 10:00（若電腦未開機則在下次開機後補跑）'),
        ('執行程式',  'python.exe  togaf_checker\\main.py  --once'),
        ('執行帳號',  '目前登入使用者（Interactive，無需系統管理員）'),
        ('執行期限',  '最長 1 小時（通常 < 5 秒）'),
    ]
    for i, (k, v) in enumerate(details, start=1):
        _shd(tb.rows[i].cells[0], F_LGRAY)
        _cell(tb.rows[i].cells[0], k, bold=True, size=10)
        _cell(tb.rows[i].cells[1], v, size=10)
    _shd(tb.rows[0].cells[0], F_NAVY)
    _shd(tb.rows[0].cells[1], F_NAVY)
    _cell(tb.rows[0].cells[0], '設定項目', bold=True, color=C_WHITE, size=10)
    _cell(tb.rows[0].cells[1], '值',       bold=True, color=C_WHITE, size=10)

    _p(doc)

    _h(doc, '4.6  CLI 參數說明', 2)
    tb2 = _tbl(doc, 6, 3, [3.5, 4, 9])
    _hdr(tb2, ['參數', '預設值', '說明'])
    args = [
        ('--root <path>',  '<togaf_checker> 上一層', '要掃描的專案根目錄'),
        ('--out  <path>',  'docs/generated/',        '報告輸出目錄（自動建立）'),
        ('--time HH:MM',   '10:00',                  'Daemon 模式每日執行時間'),
        ('--daemon',       '（開關，預設關）',        '常駐模式，每日在 --time 時執行'),
        ('--once',         '（開關，預設開）',        '執行一次後退出（預設行為）'),
    ]
    for i, (a, b, c) in enumerate(args, start=1):
        if i % 2 == 0:
            for cell in tb2.rows[i].cells:
                _shd(cell, F_LGRAY)
        _cell(tb2.rows[i].cells[0], a, bold=True, size=10)
        p = tb2.rows[i].cells[1].paragraphs[0]
        p.space_before = Pt(2); p.space_after = Pt(2)
        run = p.add_run(b)
        run.font.size = Pt(9); run.font.name = 'Courier New'
        _cell(tb2.rows[i].cells[2], c, size=10)

    _p(doc)

    _h(doc, '4.7  輸出報告說明', 2)
    _p(doc, '每次執行產生一份獨立的 Word 報告，檔名格式：', size=11)
    _code(doc, '台水TOGAF合規報告_20260518_1355.docx\n'
               '                 ^^^^^^^^ ^^^^^\n'
               '                 日期     時間（HHMM）')
    _p(doc, '報告結構：', size=11)
    tb3 = _tbl(doc, 6, 2, [3.5, 13])
    _shd(tb3.rows[0].cells[0], F_NAVY); _shd(tb3.rows[0].cells[1], F_NAVY)
    _cell(tb3.rows[0].cells[0], '章節', bold=True, color=C_WHITE, size=10)
    _cell(tb3.rows[0].cells[1], '內容', bold=True, color=C_WHITE, size=10)
    sections = [
        ('封面',          '整體合規分數、掃描目錄（完整路徑）、掃描時間戳記'),
        ('執行摘要',      '8 項準則 Pass/Fail 狀態表、各準則問題數與分數'),
        ('準則詳細頁 ×8', '每項準則：彩色問題清單（含完整檔案路徑、行號、程式碼片段）、改善建議'),
        ('附錄',          '高/中/低風險問題數量統計表'),
        ('頁首/頁尾',     '每頁顯示報告標題 + 生成時間戳記 + 頁碼'),
    ]
    for i, (s, d) in enumerate(sections, start=1):
        if i % 2 == 0:
            for c in tb3.rows[i].cells:
                _shd(c, F_LGRAY)
        _cell(tb3.rows[i].cells[0], s, bold=True, size=10)
        _cell(tb3.rows[i].cells[1], d, size=10)

    _p(doc)

    _h(doc, '4.8  環境變數覆寫（進階）', 2)
    _p(doc, '可用環境變數取代 CLI 參數，適合在 Task Scheduler action 中直接設定：', size=11)
    tb4 = _tbl(doc, 5, 3, [5, 4, 7.5])
    _hdr(tb4, ['環境變數', '預設值', '說明'])
    envs = [
        ('TOGAF_PROJECT_ROOT', '<腳本上一層目錄>', '掃描目標根目錄'),
        ('TOGAF_OUTPUT_DIR',   'docs/generated/',  '報告輸出目錄'),
        ('TOGAF_FILENAME',     '台水TOGAF合規報告_{datetime}.docx', '檔名樣板'),
        ('TOGAF_TIME',         '10:00',             'Daemon 執行時間'),
    ]
    for i, (e, d, s) in enumerate(envs, start=1):
        if i % 2 == 0:
            for c in tb4.rows[i].cells:
                _shd(c, F_LGRAY)
        _cell(tb4.rows[i].cells[0], e, bold=True, size=9)
        p = tb4.rows[i].cells[1].paragraphs[0]
        p.space_before = Pt(2); p.space_after = Pt(2)
        run = p.add_run(d); run.font.size = Pt(9); run.font.name = 'Courier New'
        _cell(tb4.rows[i].cells[2], s, size=10)

    _p(doc)
    _page_break(doc)


# ── Chapter 5: Design Document ────────────────────────────────────────────────

def _ch5_design(doc):
    _h(doc, '5. 設計文件', 1)

    # ── 5.1 Architecture ──────────────────────────────────────────────────────
    _h(doc, '5.1  系統架構總覽', 2)
    _p(doc,
       '系統由三個主要模組組成，以單向資料流運行：原始碼掃描 → 規則引擎分析 → '
       'Word 報告產生。排程由作業系統負責觸發，工具本身無需常駐 daemon 程序。',
       size=11)
    _p(doc)

    _code(doc,
          '┌─────────────────────────────────────────────────────────────────┐\n'
          '│                   TOGAF Checker — 系統架構                      │\n'
          '├──────────────┬──────────────────────────────┬───────────────────┤\n'
          '│  【觸發層】   │     【分析層】                │  【輸出層】       │\n'
          '│              │                              │                   │\n'
          '│  Windows     │  main.py                    │  report_builder   │\n'
          '│  Task        │  ┌──────────────────┐       │  .py              │\n'
          '│  Scheduler   │  │   checker.py     │       │                   │\n'
          '│  每天 10:00  ├─►│  P15 Security    ├──────►│  .docx            │\n'
          '│              │  │  P16 Tech Indep  │       │  ・封面            │\n'
          '│  或手動：    │  │  P20 Diversity   │       │  ・執行摘要        │\n'
          '│  python      │  │  P21 Interop     │       │  ・8 準則詳細頁   │\n'
          '│  main.py     │  │  P6  Service     │       │  ・附錄統計        │\n'
          '│              │  │  P5  CommonUse   │       │                   │\n'
          '│              │  │  P4  BizCont     │       │  docs/generated/  │\n'
          '│              │  │  P7  Compliance  │       │  台水TOGAF合規     │\n'
          '│              │  └──────────────────┘       │  報告_日期_時間   │\n'
          '│              │         ▲                   │  .docx            │\n'
          '│              │         │ rglob scan        │                   │\n'
          '│              │  ┌──────┴───────┐           │                   │\n'
          '│              │  │ .java files  │           │                   │\n'
          '│              │  │ .jsp  files  │           │                   │\n'
          '│              │  │ .properties  │           │                   │\n'
          '│              │  └──────────────┘           │                   │\n'
          '└──────────────┴──────────────────────────────┴───────────────────┘'
    )
    _p(doc)

    # ── 5.2 Module design ─────────────────────────────────────────────────────
    _h(doc, '5.2  模組設計', 2)
    tb = _tbl(doc, 8, 4, [3.8, 3.2, 2.5, 7.0])
    _hdr(tb, ['檔案', '主要類別/函式', '輸入', '輸出/職責'])
    mods = [
        ('principles_config.py', 'PRINCIPLES\nRECOMMENDATIONS\nget_spec()\nget_recommendations()',
         '（靜態設定，無輸入）',
         '★ 唯一的原則定義檔。定義準則 ID/名稱/優先度/說明；'
         '提供改善建議清單；checker.py 與 report_builder.py 皆從此匯入。'),
        ('checker.py',        'run_all_checks()\n8× check_pXX()',
         'project_root: str', 'ProjectCheckResult（含 N 個 PrincipleResult）'),
        ('report_builder.py', 'build_report()',
         'ProjectCheckResult\nout_path: str', '寫出 .docx；封面/摘要/N 節/附錄'),
        ('main.py',           'main()\ngenerate()\nrun_daemon()',
         'CLI 參數 / 環境變數', '協調 checker + report_builder；處理原子寫入與日誌輪替'),
        ('setup_scheduler.ps1','（PowerShell 腳本）',
         '無（互動式）', '在 Windows Task Scheduler 中新增/更新每日工作'),
        ('setup_scheduler.bat','（批次啟動器）',
         '無', '以正確編碼呼叫 .ps1，提供雙擊入口'),
        ('generate_guide.py', '（本腳本）',
         '無', '產生本使用者指南 Word 文件'),
    ]
    for i, (f, c, inp, out) in enumerate(mods, start=1):
        if i % 2 == 0:
            for cell in tb.rows[i].cells:
                _shd(cell, F_LGRAY)
        # Highlight principles_config.py in gold
        if i == 1:
            for cell in tb.rows[i].cells:
                _shd(cell, F_AMBER)
        p = tb.rows[i].cells[0].paragraphs[0]
        p.space_before = Pt(2); p.space_after = Pt(2)
        run = p.add_run(f); run.font.size = Pt(9); run.font.bold = True
        run.font.name = 'Courier New'
        _cell(tb.rows[i].cells[1], c, size=9)
        _cell(tb.rows[i].cells[2], inp, size=9, italic=True)
        _cell(tb.rows[i].cells[3], out, size=10)

    _p(doc)

    # ── 5.3 Data model ────────────────────────────────────────────────────────
    _h(doc, '5.3  資料模型', 2)
    _code(doc,
          'ProjectCheckResult\n'
          '  root: str                  # 專案根目錄完整路徑\n'
          '  generated_at: str          # "2026-05-18 13:55:08"\n'
          '  total_files_scanned: int\n'
          '  results: List[PrincipleResult]\n'
          '\n'
          'PrincipleResult\n'
          '  principle_id: str          # "P15"\n'
          '  principle_name: str        # "Data Security"\n'
          '  priority: str              # "最高" | "高"\n'
          '  checks_run: int            # 本準則執行的檢查總數\n'
          '  checks_passed: int         # 通過的檢查數\n'
          '  findings: List[Finding]    # 違規清單\n'
          '  passed: bool               # 無 HIGH/MEDIUM finding = True\n'
          '  score: int                 # checks_passed / checks_run × 100\n'
          '\n'
          'Finding\n'
          '  filepath: str              # 相對於 project_root 的完整路徑\n'
          '  line_no: int               # 行號（0 = 檔案層級）\n'
          '  evidence: str              # 原始程式碼片段（≤120 字元）\n'
          '  severity: str              # "HIGH" | "MEDIUM" | "LOW"\n'
          '  rule: str                  # 違規規則說明'
    )
    _p(doc)

    # ── 5.4 Execution flow ────────────────────────────────────────────────────
    _h(doc, '5.4  執行流程', 2)
    _code(doc,
          'Task Scheduler 觸發\n'
          '    │\n'
          '    ▼\n'
          'main.py  generate(project_root, output_dir)\n'
          '    │\n'
          '    ├─► checker.run_all_checks(root)\n'
          '    │       │\n'
          '    │       ├─► collect all .java files  (pathlib.rglob, skip _skip())\n'
          '    │       ├─► collect all .jsp  files  (pathlib.rglob)\n'
          '    │       ├─► collect all .properties  (pathlib.rglob)\n'
          '    │       │\n'
          '    │       ├─► check_p15()  → regex scan each line of each file\n'
          '    │       ├─► check_p16()  → regex scan + interface ratio\n'
          '    │       ├─► check_p20()  → multi-library detection across all files\n'
          '    │       ├─► check_p21()  → IP/Content-Type/API-doc checks\n'
          '    │       ├─► check_p6()   → layer detection + method count\n'
          '    │       ├─► check_p5()   → class-name duplicate map\n'
          '    │       ├─► check_p4()   → empty-catch + broad-catch + null\n'
          '    │       └─► check_p7()   → PII in logs + audit log + masking\n'
          '    │\n'
          '    ├─► report_builder.build_report(result, tmp_path)\n'
          '    │       │\n'
          '    │       ├─► cover page  (score gauge, timestamp, full project path)\n'
          '    │       ├─► summary table\n'
          '    │       ├─► 8 × principle section  (findings with relative paths)\n'
          '    │       └─► appendix statistics\n'
          '    │\n'
          '    └─► atomic rename:  .tmp.docx → 台水TOGAF合規報告_YYYYMMDD_HHMM.docx'
    )
    _p(doc)

    # ── 5.5 Skip / filter logic ───────────────────────────────────────────────
    _h(doc, '5.5  檔案過濾邏輯', 2)
    _p(doc,
       '掃描時自動略過備份/建置產出/工具目錄，確保報告反映真實原始碼：',
       size=11)
    tb2 = _tbl(doc, 3, 3, [4, 4, 8.5])
    _hdr(tb2, ['過濾類型', '觸發條件', '範例'])
    filters = [
        ('目錄黑名單',
         'SKIP_DIRS 精確比對路徑元件',
         '.git  .claude  target  build  __pycache__\nnode_modules  .tmp_ccms  update  archive'),
        ('備份/複製資料夾',
         'COPY_FOLDER_MARKERS 子字串比對（大小寫不分）',
         'waterImportTimer - ciopy\nwaterHrSync - 複製\nwaterCusSSO_backup'),
    ]
    for i, (a, b, c) in enumerate(filters, start=1):
        _cell(tb2.rows[i].cells[0], a, bold=True, size=10)
        _cell(tb2.rows[i].cells[1], b, size=10)
        p = tb2.rows[i].cells[2].paragraphs[0]
        p.space_before = Pt(2); p.space_after = Pt(2)
        run = p.add_run(c); run.font.size = Pt(9); run.font.name = 'Courier New'

    _p(doc)

    # ── 5.6 Atomic write ──────────────────────────────────────────────────────
    _h(doc, '5.6  原子寫入設計（防止 Word 鎖定衝突）', 2)
    _p(doc,
       '報告先寫入 .tmp.docx 暫存檔，完成後再以 Path.rename() 替換目標檔案。'
       '若目標檔案正被 Microsoft Word 開啟而鎖定，則自動在檔名加上時間戳記另存：',
       size=11)
    _code(doc,
          '# main.py  generate()\n'
          'tmp_path = out_path.with_suffix(".tmp.docx")\n'
          'build_report(result, str(tmp_path))      # 寫入暫存檔\n'
          '\n'
          'if out_path.exists():\n'
          '    try:\n'
          '        out_path.unlink()                # 刪除舊版\n'
          '    except PermissionError:\n'
          '        ts  = datetime.now().strftime("%H%M%S")\n'
          '        out_path = out_path.with_stem(out_path.stem + f"_{ts}")\n'
          '\n'
          'tmp_path.rename(out_path)                # 原子替換'
    )
    _p(doc)

    # ── 5.7 principles_config.py — single source of truth ────────────────────
    _h(doc, '5.7  principles_config.py — 原則設定檔設計', 2)
    _p(doc,
       '新增 principles_config.py 作為所有準則定義的「唯一真相來源（Single Source of Truth）」。'
       'checker.py 與 report_builder.py 皆從此檔匯入，避免同一份資料分散在多個地方維護。',
       size=11)
    _p(doc)

    # Dependency diagram
    _code(doc,
          '                  ┌──────────────────────────┐\n'
          '                  │   principles_config.py   │\n'
          '                  │  ★ 唯一修改入口           │\n'
          '                  │                          │\n'
          '                  │  PRINCIPLES[]            │\n'
          '                  │  RECOMMENDATIONS{}       │\n'
          '                  │  get_spec(id)            │\n'
          '                  │  get_recommendations(id) │\n'
          '                  └──────────┬───────────────┘\n'
          '                   匯入       │      匯入\n'
          '          ┌────────────────────┼────────────────────┐\n'
          '          ▼                    ▼                    ▼\n'
          '    checker.py          report_builder.py    （未來擴充）\n'
          '    讀取準則的           讀取改善建議\n'
          '    id/name/priority     get_recommendations()'
    )
    _p(doc)

    # Data structure
    _h(doc, '5.7.1  資料結構', 3)
    _code(doc,
          '# PrincipleSpec — 每條準則的 metadata\n'
          '@dataclass\n'
          'class PrincipleSpec:\n'
          '    id:             str   # "P15"\n'
          '    name:           str   # "Data Security"\n'
          '    priority:       str   # "最高" | "高" | "中" | "低"\n'
          '    description:    str   # 一行說明，顯示於報告準則標題下方\n'
          '    checks_summary: str   # 逗號分隔的檢查重點，顯示於準則參考表\n'
          '\n'
          '# PRINCIPLES — 有序清單，決定報告章節順序\n'
          'PRINCIPLES: List[PrincipleSpec] = [\n'
          '    PrincipleSpec(id="P15", name="Data Security", priority="最高", ...),\n'
          '    PrincipleSpec(id="P16", name="Technology Independence", ...),\n'
          '    ...  # 依優先度排序\n'
          ']\n'
          '\n'
          '# RECOMMENDATIONS — 每條準則的改善建議清單\n'
          'RECOMMENDATIONS: dict[str, List[str]] = {\n'
          '    "P15": ["使用設定檔管理密碼，禁止硬編碼", "採用 PreparedStatement ...", ...],\n'
          '    "P16": [...],\n'
          '    ...\n'
          '}'
    )
    _p(doc)

    # How to maintain table
    _h(doc, '5.7.2  維護指引', 3)
    _p(doc, '根據不同情境，只需修改對應的位置：', size=11)

    tb3 = _tbl(doc, 6, 3, [4.5, 4.5, 7.5])
    _hdr(tb3, ['情境', '修改位置', '具體操作'])
    ops = [
        ('改準則名稱 / 優先度 / 說明',
         'principles_config.py\n只改這一個檔案',
         '修改 PRINCIPLES 清單中對應 PrincipleSpec 的欄位值'),
        ('改改善建議（報告中的條列）',
         'principles_config.py\n只改這一個檔案',
         '修改 RECOMMENDATIONS[\"PXX\"] 的字串清單'),
        ('新增一條原則',
         '① principles_config.py\n② checker.py',
         '① 在 PRINCIPLES 加 PrincipleSpec；在 RECOMMENDATIONS 加建議清單\n'
         '② 新增 check_pXX() 函式，並加入 CHECKERS 清單'),
        ('刪除一條原則',
         '① principles_config.py\n② checker.py',
         '① 從 PRINCIPLES 與 RECOMMENDATIONS 刪除該 ID\n'
         '② 刪除對應的 check_pXX() 函式，並從 CHECKERS 移除'),
        ('改偵測規則（regex）',
         'checker.py\n只改這一個檔案',
         '修改 check_pXX() 函式內的 regex 常數或判斷邏輯'),
    ]
    op_fills = [F_GREEN, F_GREEN, F_AMBER, F_AMBER, F_BLUE]
    for i, (sit, loc, act) in enumerate(ops, start=1):
        _shd(tb3.rows[i].cells[0], op_fills[i - 1])
        _cell(tb3.rows[i].cells[0], sit, bold=True, color=C_WHITE, size=10)
        # Location cell — monospace
        loc_cell = tb3.rows[i].cells[1]
        loc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = loc_cell.paragraphs[0]
        lp.space_before = Pt(3); lp.space_after = Pt(3)
        for j, line in enumerate(loc.split('\n')):
            if j > 0:
                lp = loc_cell.add_paragraph()
                lp.space_before = Pt(1); lp.space_after = Pt(1)
            run = lp.add_run(line)
            run.font.size = Pt(9)
            run.font.bold = (j == 0)
            run.font.name = 'Courier New' if j == 0 else None
        # Action cell — multi-line
        act_cell = tb3.rows[i].cells[2]
        act_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        ap = act_cell.paragraphs[0]
        ap.space_before = Pt(3); ap.space_after = Pt(3)
        for j, line in enumerate(act.split('\n')):
            if j > 0:
                ap = act_cell.add_paragraph()
                ap.space_before = Pt(1); ap.space_after = Pt(1)
            run = ap.add_run(line)
            run.font.size = Pt(9.5)

    _p(doc)

    # Validation steps
    _h(doc, '5.7.3  修改後驗證步驟', 3)
    _numbered(doc, '確認語法正確：python principles_config.py')
    _numbered(doc, '確認掃描結果符合預期：python checker.py D:\\project\\water')
    _numbered(doc, '產生報告確認排版：python main.py')
    _numbered(doc, '（可選）重新產生本使用者指南：python generate_guide.py')

    _p(doc)
    _page_break(doc)


# ── Chapter 6: TOGAF Principles reference ────────────────────────────────────

def _ch6_togaf(doc):
    _h(doc, '6. TOGAF 準則參考表', 1)
    _p(doc,
       '本工具對應 TOGAF 框架中 8 項與軟體開發直接相關的準則，'
       '優先度由業務與系統架構風險共同決定。',
       size=11)
    _p(doc)

    tb = _tbl(doc, 9, 5, [1.5, 1.5, 3.5, 4.0, 6.0])
    _hdr(tb, ['優先度', '準則 ID', '準則名稱', '核心要求', '本工具檢查重點'])
    principles = [
        ('🔴 最高', '#15', 'Data Security',
         '加密、認證、授權、防注入攻擊',
         '硬編碼密碼、SQL 注入、Token 繞過、CORS 萬用字元'),
        ('🔴 最高', '#16', 'Technology Independence',
         '抽象層、依賴注入、避免 Vendor Lock-in',
         '直接 new *Impl()、廠商特定 import、介面比例'),
        ('🔴 最高', '#20', 'Control Technical Diversity',
         '統一技術棧、遵守語言/框架規範',
         '多 JSON 函式庫、多 Log 框架、多 HTTP 用戶端並存'),
        ('🔴 最高', '#21', 'Interoperability',
         '標準 API、訊息格式（JSON/XML）、協議',
         '硬編碼 IP、缺 Content-Type、缺 Swagger 文件'),
        ('🟠 高',   '#6',  'Service Orientation',
         '微服務/模組化設計、鬆耦合',
         '缺分層（Model/DAO/Service/Action）、過胖類別'),
        ('🟠 高',   '#5',  'Common Use Applications',
         '共用元件、禁止重複造輪子',
         '跨模組重複類別定義、缺 common/util 套件'),
        ('🟠 高',   '#4',  'Business Continuity',
         '錯誤處理、重試機制、熔斷器',
         '空 catch、無日誌的寬泛 catch、缺 null 防禦'),
        ('🟠 高',   '#7',  'Compliance with Law',
         '個資脫敏、稽核日誌',
         '個資寫入日誌、缺脫敏處理、缺稽核日誌'),
    ]
    prio_fills = {'🔴 最高': F_RED, '🟠 高': F_AMBER}
    for i, (prio, pid, pname, req, checks) in enumerate(principles, start=1):
        fill = prio_fills[prio]
        _shd(tb.rows[i].cells[0], fill)
        _cell(tb.rows[i].cells[0], prio, bold=True, color=C_WHITE,
              size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(tb.rows[i].cells[1], pid, bold=True, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(tb.rows[i].cells[2], pname, bold=True, size=10)
        _cell(tb.rows[i].cells[3], req, size=10)
        _cell(tb.rows[i].cells[4], checks, size=10)

    _p(doc)
    _p(doc,
       '準則全文來源：TOGAF Principles.docx（台水資訊整合系統專案文件）',
       size=10, italic=True, color=C_LGRAY)

    _p(doc)
    _page_break(doc)


# ── Chapter 7: FAQ & Troubleshooting ─────────────────────────────────────────

def _ch7_faq(doc):
    _h(doc, '7. 常見問題（FAQ）', 1)

    faqs = [
        ('Q1：執行報告只花了 5 秒，真的有掃描所有檔案嗎？',
         '是的。靜態 regex 掃描比執行程式快得多。32 個 .java/.jsp 檔案大約 '
         '3–10 秒，即使是數百個檔案也通常在 30 秒內完成。\n'
         '掃描的檔案數量會記錄在報告封面（「掃描檔案總數：XX 個」）。'),
        ('Q2：每次報告分數都是 100%，但有 FAIL 準則，這正確嗎？',
         '正確。分數計算基礎是「行級別檢查通過數 / 行級別檢查總數」。\n'
         '若一個檔案有 500 行，而只有 2 行觸發 HIGH 規則，分數仍達 99.6% ≈ 100%。\n'
         '準則的 Pass/Fail 判斷取決於「是否有 HIGH 或 MEDIUM 等級的 finding」，\n'
         '這才是需要關注的指標，而非分數百分比本身。'),
        ('Q3：排程工作設定後為何沒有在 10:00 執行？',
         '常見原因：\n'
         '1. 電腦在 10:00 時關機或睡眠 → 啟用「StartWhenAvailable」（已在 PS1 設定）\n'
         '2. 工作狀態為 Disabled → 開啟「工作排程器」確認狀態為 Ready\n'
         '3. python.exe 路徑不在 PATH → 在 cmd 執行 where python 確認'),
        ('Q4：報告說「缺少稽核日誌」，但我的程式有寫 log，為什麼？',
         '工具只偵測包含 audit、AuditLog、稽核 關鍵字的日誌呼叫。\n'
         '一般的 logger.info() 不算稽核日誌。\n'
         '建議建立專用的 AuditLogService 類別，明確標記關鍵操作的稽核記錄。'),
        ('Q5：如何新增自訂的檢查規則（regex pattern）？',
         '在 checker.py 中對應的 check_pXX() 函式內：\n'
         '1. 新增正規表達式 pattern（模組頂層常數或函式內 local）\n'
         '2. 在迴圈中對每行執行 pattern.search(line)\n'
         '3. 違規時呼叫 result.add_finding(Finding(rel, ln, line, "HIGH", "規則說明"))\n'
         '無需修改 main.py、report_builder.py 或 principles_config.py。'),
        ('Q6：如何新增一條全新的 TOGAF 原則？',
         '需要修改兩個檔案（詳見 5.7.2 維護指引）：\n'
         '① principles_config.py\n'
         '   - 在 PRINCIPLES 清單加入新的 PrincipleSpec(id="P22", name="...", priority="高", ...)\n'
         '   - 在 RECOMMENDATIONS 加入對應建議清單\n'
         '② checker.py\n'
         '   - 新增 check_p22(root) 函式，回傳 PrincipleResult\n'
         '   - 在 CHECKERS 清單末尾加入 check_p22\n'
         '驗證：python checker.py D:\\project\\water'),
        ('Q7：可以在 Linux 或 Mac 上執行嗎？',
         '可以。除了 setup_scheduler.ps1（Windows 專屬），所有 Python 腳本均跨平台。\n'
         'Linux/Mac 排程：使用 cron\n'
         '  0 10 * * * /usr/bin/python3 /path/to/togaf_checker/main.py'),
    ]

    for q, a in faqs:
        tb = _tbl(doc, 2, 1, [16.5])
        _shd(tb.rows[0].cells[0], F_NAVY)
        _cell(tb.rows[0].cells[0], q, bold=True, color=C_WHITE, size=10)
        # Answer cell with code formatting preserved
        ans_cell = tb.rows[1].cells[0]
        ans_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = ans_cell.paragraphs[0]
        p.space_before = Pt(4)
        p.space_after  = Pt(4)
        for line in a.split('\n'):
            if p.text:
                p = ans_cell.add_paragraph()
                p.space_before = Pt(1)
                p.space_after  = Pt(1)
            is_code = line.startswith('  ') or line.startswith('0 10')
            run = p.add_run(line)
            run.font.size = Pt(9 if is_code else 10)
            if is_code:
                run.font.name = 'Courier New'
        _p(doc)

    _page_break(doc)


# ── Chapter 8: Version history ────────────────────────────────────────────────

def _ch8_changelog(doc):
    _h(doc, '8. 版本紀錄', 1)
    tb = _tbl(doc, 5, 4, [2, 3, 5, 7])
    _hdr(tb, ['版本', '日期', '異動項目', '說明'])
    changes = [
        ('1.0.0', '2026-05-18',
         '初始版本',
         '實作 8 項 TOGAF 準則規則引擎；生成彩色 Word 報告；'
         'Windows Task Scheduler 排程設定'),
        ('1.0.1', '2026-05-18',
         '報告改善',
         '檔名加入時間戳記（YYYYMMDD_HHMM）；'
         '問題檔案路徑改為相對專案根目錄的完整路徑；'
         '頁首加入生成時間戳記'),
        ('1.0.2', '2026-05-18',
         '架構重構：principles_config.py',
         '新增 principles_config.py 作為原則定義的唯一真相來源；'
         'checker.py 與 report_builder.py 改從此檔匯入準則 metadata 與改善建議；'
         '新增使用者指南 5.7 維護指引章節'),
        ('（計劃）', '—',
         '—',
         '增加 XML / Groovy / Kotlin 檔案支援；增加 HTML 格式報告輸出'),
    ]
    for i, (v, d, c, s) in enumerate(changes, start=1):
        if i % 2 == 0:
            for cell in tb.rows[i].cells:
                _shd(cell, F_LGRAY)
        # Highlight current version
        if v == '1.0.2':
            for cell in tb.rows[i].cells:
                _shd(cell, F_TEAL)
        _cell(tb.rows[i].cells[0], v, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
              color=C_WHITE if v == '1.0.2' else C_BLACK)
        _cell(tb.rows[i].cells[1], d, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
              color=C_WHITE if v == '1.0.2' else C_BLACK)
        _cell(tb.rows[i].cells[2], c, bold=True, size=10,
              color=C_WHITE if v == '1.0.2' else C_BLACK)
        _cell(tb.rows[i].cells[3], s, size=10,
              color=C_WHITE if v == '1.0.2' else C_BLACK)

    _p(doc)
    _p(doc,
       f'本文件由 generate_guide.py 自動產生  ·  '
       f'工具版本 1.0.2  ·  生成時間：{TODAY}',
       size=9, color=C_LGRAY, italic=True,
       align=WD_ALIGN_PARAGRAPH.CENTER)


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════

def build_guide(out_path: str):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    _add_header_footer(
        doc,
        f'TOGAF Checker 使用者指南  ·  v1.0.2  ·  {TODAY}'
    )

    _cover(doc)
    _ch1_intro(doc)
    _ch2_ai(doc)
    _ch3_stack(doc)
    _ch4_howto(doc)
    _ch5_design(doc)
    _ch6_togaf(doc)
    _ch7_faq(doc)
    _ch8_changelog(doc)

    doc.save(out_path)
    return out_path


def main():
    parser = argparse.ArgumentParser(
        description='Generate TOGAF Checker User Guide in Word format.')
    parser.add_argument('--out', default=str(Path(__file__).parent.parent / 'docs'),
                        help='Output directory (default: ../docs/)')
    args = parser.parse_args()

    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = str(out_dir / 'togaf_checker_user_guide.docx')

    print(f'Generating: {out_path} ...')
    build_guide(out_path)
    size_kb = Path(out_path).stat().st_size // 1024
    print(f'[OK] {out_path}  ({size_kb} KB)')


if __name__ == '__main__':
    main()
